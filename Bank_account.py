import db
import config
from datetime import datetime
import time
import random as rd
import smtplib
from email.message import EmailMessage
import os
import docx

BANK_EMAIL = config.BANK_EMAIL
BANK_EMAIL_PASSWORD = config.BANK_EMAIL_PASSWORD


class Account():
    def __init__(self, balance):
        self.balance = balance

    def show_account_balance(self,):
        self.balance = Amount.check_amount()
        print(f"Your account balance is {self.balance}")


class Changepassword():
    def __init__(self, email, new_pin, code):
        self.email = email
        self.new_pin = new_pin
        self.servermail = BANK_EMAIL
        self.serverpassword = BANK_EMAIL_PASSWORD
        self.confidential_code = code

    def confirm_email(self):
        msg = EmailMessage()
        msg['Subject'] = 'Confidential'
        msg['From'] = self.servermail
        msg['To'] = self.email
        msg['Reply-To'] = "no-reply@gmail.com"
        msg.set_content(
            f"One-Time Password Reset Code — Action Required \n"
            f"We received a request to reset your password. Use the code below to proceed: "
            f"Reset Code: {self.confidential_code} "
            f"This code expires in 10 minutes and is for single use only. "
            f"Never share it with anyone. If you did not request this, please ignore this message. "
            f"Your password remains unchanged"
        )
        try:
            with smtplib.SMTP("smtp.gmail.com", 587) as server:
                server.ehlo()
                server.starttls()
                server.login(self.servermail, self.serverpassword)
                server.send_message(msg)
            print("Email sent successfully!")
        except smtplib.SMTPAuthenticationError:
            print(" ERROR: Email authentication failed. Check your email credentials and app password.")
        except smtplib.SMTPException as e:
            print(f" ERROR: Email sending failed. {str(e)}")
        except Exception as e:
            print(f" ERROR: {str(e)}")

    def change(self):
        conn = db.get_connection()
        cursor = conn.cursor()
        cursor.execute(
            "UPDATE Customer_services SET Customer_password = %s WHERE Customer_Email = %s",
            (self.new_pin, self.email)
        )
        conn.commit()
        conn.close()
        print("Your password was updated successfully!")


class Amount():
    def __init__(self, account: str):
        self.amount = int()
        self.account = account

    def check_amount(self):
        conn = db.get_connection()
        cur = conn.cursor()
        cur.execute("SELECT * FROM Account WHERE Account_id = %s", (self.account,))
        amount = cur.fetchone()
        conn.close()
        if amount is None:
            return "No Transaction recorded"
        return amount[0]


class Transaction():
    def __init__(self, amount: int, recipient_account, sender_account=None):
        self.amount = amount
        self.transaction_id = int()
        self.account_id = recipient_account
        self.sender_account = sender_account

    def check(self):
        account_to_check = self.sender_account if self.sender_account else self.account_id
        conn = db.get_connection()
        cursor = conn.cursor()
        cursor.execute(
            "SELECT Current_amount FROM Account WHERE Account_id = %s",
            (str(account_to_check),)
        )
        amount = cursor.fetchone()
        conn.close()
        if amount is None:
            print("Account not found.")
            return False
        current = float(amount[0])
        if current >= self.amount:
            return True
        else:
            print("Insufficient funds. Please recharge!")
            return False

    def send(self, confirm: str) -> bool:
        if confirm:
            sender = self.sender_account if self.sender_account else self.account_id
            conn = db.get_connection()
            cursor = conn.cursor()
            cursor.execute(
                "UPDATE Account SET Current_amount = Current_amount - %s WHERE Account_id = %s",
                (self.amount, str(sender))
            )
            cursor.execute(
                "UPDATE Account SET Current_amount = Current_amount + %s WHERE Account_id = %s",
                (self.amount, str(self.account_id))
            )
            cursor.execute(
                "SELECT Current_amount FROM Account WHERE Account_id = %s",
                (str(sender),)
            )
            balance = cursor.fetchone()
            new_balance = balance[0] if balance is not None else None
            conn.commit()
            conn.close()
            print("Transaction was successful!")
            return new_balance
        return None

    def transaction_generator(self):
        transac_1 = rd.sample(range(1, 10), 9)
        rd.shuffle(transac_1)
        id_initial = str()
        for c in transac_1:
            id_initial += str(c)
        self.transaction_id = int(id_initial)
        return self.transaction_id

    def transaction_report(self):
        transaction_day = datetime.today()
        transaction_time = time.strftime("%H:%M:%S")
        with open("Bank_Transaction.txt", "a") as transact:
            transact.write("Transaction Initiated\n")
            transact.write(
                f"Account {self.account_id} sent GHc{self.amount:.2f} "
                f"at {transaction_time}, {transaction_day}\n"
            )
            transact.write("Transaction Successful \n")
        return transaction_day, transaction_time

    def transaction_record(self):
        transaction_day = datetime.today()
        conn = db.get_connection()
        cur = conn.cursor()
        cur.execute(
            "UPDATE Account SET Transaction_ID = %s WHERE Account_id = %s",
            (self.transaction_id, self.account_id)
        )
        cur.execute(
            "UPDATE Account SET Transaction_Date = %s WHERE Account_id = %s",
            (transaction_day, self.account_id)
        )
        conn.commit()
        conn.close()

    def report_csv(self):
        """Generate CSV report of transactions"""
        pass


def customer_message(msg: str):
    filename = "Bank_customer_report.docx"
    if not os.path.exists(filename):
        fileobject = docx.Document()
        para1 = fileobject.add_paragraph("-------------------Report----------------")
        para1.add_run("\n" + msg)
        fileobject.save(filename)
    else:
        another_fileobject = docx.Document(filename)
        apara1 = another_fileobject.add_paragraph("-------------------Report----------------")
        apara1.add_run(msg)
        another_fileobject.save(filename)


def get_account_number(email: str):
    conn = db.get_connection()
    cur = conn.cursor()
    cur.execute("SELECT * FROM Customer_services WHERE Customer_Email = %s", (email,))
    account = cur.fetchone()
    cur.execute("SELECT * FROM Account WHERE Account_id = %s", (account[8],))
    current_amount = cur.fetchone()
    conn.close()
    return account[8], int(current_amount[2])
