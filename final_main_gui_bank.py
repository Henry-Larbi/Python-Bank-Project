"""
Access Bank – Final Integrated GUI
Integrates all logic from:
  • Bank_bulid_1.py  (registration, login, password generator, send_email)
  • Bank_account.py  (Transaction, Changepassword, Amount, customer_message)
  • Bank_interface.py (registration validations, username generation)
  • Bank_Account_main.py (dashboard: send money, OTP change password, customer care)
Single self-contained file.
"""

import sys
import os
import re
import csv
import time
import random
import string
import smtplib
from datetime import datetime, timedelta
from email.message import EmailMessage

import db
import config

from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QPushButton, QLineEdit, QLabel, QDialog, QMessageBox, QComboBox,
    QTabWidget, QTableWidget, QTableWidgetItem, QSpinBox, QDoubleSpinBox,
    QTextEdit, QFormLayout, QGroupBox, QStackedWidget, QAction,
    QHeaderView, QFrame, QSizePolicy
)
from PyQt5.QtCore import Qt, pyqtSignal, QTimer
from PyQt5.QtGui import QFont, QColor, QPalette

try:
    import docx as _docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# ── Bank credentials loaded from config.py ───────────────────────────────────
BANK_EMAIL      = config.BANK_EMAIL
BANK_EMAIL_PASSWORD = config.BANK_EMAIL_PASSWORD
ADMIN_EMAIL     = config.ADMIN_EMAIL
ADMIN_PASSWORD  = config.ADMIN_PASSWORD


# ═══════════════════════════════════════════════════════════════════════════════
#  DATABASE  (from Bank_bulid_1.py + Bank_account.py)
# ═══════════════════════════════════════════════════════════════════════════════
class DB:
    @staticmethod
    def init():
        """Ensure all required tables exist — safe to call on every startup."""
        conn = db.get_connection()
        cur = conn.cursor()
        cur.execute('''
            CREATE TABLE IF NOT EXISTS Customer_services(
                Customer_Name     TEXT NOT NULL,
                Customer_Age      TEXT NOT NULL,
                Customer_Gender   TEXT NOT NULL,
                Customer_Status   TEXT NOT NULL,
                Customer_Location TEXT NOT NULL,
                Customer_Phone    TEXT NOT NULL,
                Customer_Email    TEXT NOT NULL UNIQUE PRIMARY KEY,
                Customer_password TEXT NOT NULL,
                Customer_ID       INT  NOT NULL UNIQUE
            )''')
        cur.execute('''
            CREATE TABLE IF NOT EXISTS Account(
                Account_id       TEXT PRIMARY KEY,
                Customer_ID      INT,
                Current_amount   REAL DEFAULT 0,
                Transaction_Date TEXT,
                Transaction_ID   TEXT
            )''')
        cur.execute('''
            CREATE TABLE IF NOT EXISTS Transactions(
                Transaction_ID   TEXT PRIMARY KEY,
                From_Account     TEXT,
                To_Account       TEXT,
                Amount           REAL,
                Transaction_Date TEXT,
                Status           TEXT
            )''')
        cur.execute('''
            CREATE TABLE IF NOT EXISTS Loans(
                Loan_ID          TEXT PRIMARY KEY,
                Account_ID       TEXT,
                Amount           REAL,
                Interest_Rate    REAL,
                Months           INTEGER,
                Monthly_Payment  REAL,
                Start_Date       TEXT,
                Status           TEXT DEFAULT 'Active',
                Amount_Repaid    REAL DEFAULT 0
            )''')
        cur.execute('''
            CREATE TABLE IF NOT EXISTS FixedDeposits(
                FD_ID          TEXT PRIMARY KEY,
                Account_ID     TEXT,
                Amount         REAL,
                Interest_Rate  REAL,
                Days           INTEGER,
                Start_Date     TEXT,
                Maturity_Date  TEXT,
                Expected_Return REAL,
                Status         TEXT DEFAULT 'Active'
            )''')
        cur.execute('''
            CREATE TABLE IF NOT EXISTS BillPayments(
                Payment_ID    TEXT PRIMARY KEY,
                Account_ID    TEXT,
                Biller        TEXT,
                Amount        REAL,
                Payment_Date  TEXT,
                Status        TEXT DEFAULT 'Paid'
            )''')
        cur.execute('''
            CREATE TABLE IF NOT EXISTS RecurringPayments(
                Schedule_ID   TEXT PRIMARY KEY,
                Account_ID    TEXT,
                Recipient_ID  TEXT,
                Amount        REAL,
                Frequency     TEXT,
                Next_Run_Date TEXT,
                Status        TEXT DEFAULT 'Active',
                Description   TEXT
            )''')
        cur.execute('''
            CREATE TABLE IF NOT EXISTS Notifications(
                Notif_ID     TEXT PRIMARY KEY,
                Account_ID   TEXT,
                Message      TEXT,
                Type         TEXT,
                Created_Date TEXT,
                Is_Read      INTEGER DEFAULT 0
            )''')
        cur.execute('''
            CREATE TABLE IF NOT EXISTS FrozenAccounts(
                Account_ID  TEXT PRIMARY KEY,
                Frozen_Date TEXT,
                Reason      TEXT
            )''')
        conn.commit()
        conn.close()

    @staticmethod
    def get_customer(email):
        conn = db.get_connection()
        cur = conn.cursor()
        cur.execute("SELECT * FROM Customer_services WHERE Customer_Email = %s", (email,))
        row = cur.fetchone()
        conn.close()
        return row

    @staticmethod
    def get_account(customer_id):
        conn = db.get_connection()
        cur = conn.cursor()
        cur.execute("SELECT * FROM Account WHERE Customer_ID = %s", (customer_id,))
        row = cur.fetchone()
        conn.close()
        return row

    @staticmethod
    def get_transactions(account_id):
        conn = db.get_connection()
        cur = conn.cursor()
        cur.execute('''
            SELECT * FROM Transactions
            WHERE From_Account = %s OR To_Account = %s
            ORDER BY Transaction_Date DESC
        ''', (str(account_id), str(account_id)))
        rows = cur.fetchall()
        conn.close()
        return rows

    @staticmethod
    def login(email, password):
        """Return True if credentials match (mirrors Bank_bulid_1.check())."""
        conn = db.get_connection()
        cur = conn.cursor()
        cur.execute("SELECT Customer_password FROM Customer_services WHERE Customer_Email = %s", (email,))
        data = cur.fetchone()
        conn.close()
        return data is not None and password == data[0]

    @staticmethod
    def is_duplicate(email):
        """Return True if email already registered (mirrors Sign_up_check.check())."""
        conn = db.get_connection()
        cur = conn.cursor()
        cur.execute("SELECT Customer_Email FROM Customer_services WHERE Customer_Email = %s", (email,))
        result = cur.fetchone()
        conn.close()
        return result is not None

    @staticmethod
    def register(name, age, gender, status, location, phone, email, password, account_id):
        """
        Register a new customer and create their Account row.
        Mirrors Register_Identity.register() / register_not() from Bank_bulid_1.py.
        """
        conn = db.get_connection()
        cur = conn.cursor()
        cur.execute('''
            INSERT INTO Customer_services(
                Customer_Name, Customer_Age, Customer_Gender, Customer_Status,
                Customer_Location, Customer_Phone, Customer_Email,
                Customer_password, Customer_ID)
            VALUES(%s,%s,%s,%s,%s,%s,%s,%s,%s)
        ''', (name, str(age), gender, status, location, phone, email, password, account_id))
        cur.execute('''
            INSERT INTO Account(Account_id, Customer_ID, Current_amount,
                                Transaction_Date, Transaction_ID)
            VALUES(%s,%s,%s,%s,%s)
        ''', (str(account_id), account_id, 0.0, None, None))
        conn.commit()
        conn.close()
        # Write text report (mirrors Bank_bulid_1.py report file)
        with open("Bank_JH.txt", "a") as f:
            f.write(
                f"{name}\n{age}\n{gender}\n{status}\n"
                f"{location}\n{phone}\n{email}\n{password}\n{account_id}\n---\n"
            )

    @staticmethod
    def update_password(email, new_password):
        """Mirrors Changepassword.change() from Bank_account.py."""
        conn = db.get_connection()
        cur = conn.cursor()
        cur.execute(
            "UPDATE Customer_services SET Customer_password = %s WHERE Customer_Email = %s",
            (new_password, email)
        )
        conn.commit()
        conn.close()

    @staticmethod
    def transfer(sender_id, recipient_id, amount):
        """
        Deduct from sender, credit recipient.
        Mirrors Transaction.check() + Transaction.send() from Bank_account.py.
        Returns (new_balance, error_message).
        """
        conn = db.get_connection()
        cur = conn.cursor()
        cur.execute("SELECT Current_amount FROM Account WHERE Account_id = %s", (str(sender_id),))
        row = cur.fetchone()
        if row is None:
            conn.close()
            return None, "Sender account not found."
        if float(row[0]) < amount:
            conn.close()
            return None, "Insufficient funds. Please recharge!"
        cur.execute("SELECT Account_id FROM Account WHERE Account_id = %s", (str(recipient_id),))
        if cur.fetchone() is None:
            conn.close()
            return None, "Recipient account not found."
        cur.execute(
            "UPDATE Account SET Current_amount = Current_amount - %s WHERE Account_id = %s",
            (amount, str(sender_id))
        )
        cur.execute(
            "UPDATE Account SET Current_amount = Current_amount + %s WHERE Account_id = %s",
            (amount, str(recipient_id))
        )
        cur.execute("SELECT Current_amount FROM Account WHERE Account_id = %s", (str(sender_id),))
        new_balance = cur.fetchone()[0]
        conn.commit()
        conn.close()
        return new_balance, None

    @staticmethod
    def record_transaction(txn_id, from_account, to_account, amount):
        """
        Persist a completed transfer.
        Mirrors Transaction.transaction_record() from Bank_account.py.
        """
        now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        conn = db.get_connection()
        cur = conn.cursor()
        cur.execute('''
            INSERT INTO Transactions(
                Transaction_ID, From_Account, To_Account,
                Amount, Transaction_Date, Status)
            VALUES(%s,%s,%s,%s,%s,%s)
        ''', (str(txn_id), str(from_account), str(to_account), amount, now, "Completed"))
        conn.commit()
        conn.close()
        # Also write text report (mirrors Transaction.transaction_report())
        with open("Bank_Transaction.txt", "a") as f:
            f.write(
                f"Transaction Initiated\n"
                f"Account {from_account} sent GHc{amount:.2f} at {now}\n"
                f"Transaction Successful\n---\n"
            )

    @staticmethod
    def deposit(account_id, amount):
        """Credit an account with the deposit amount. Returns (new_balance, error)."""
        conn = db.get_connection()
        cur = conn.cursor()
        cur.execute("SELECT Account_id FROM Account WHERE Account_id = %s", (str(account_id),))
        if cur.fetchone() is None:
            conn.close()
            return None, "Account not found."
        cur.execute(
            "UPDATE Account SET Current_amount = Current_amount + %s WHERE Account_id = %s",
            (amount, str(account_id))
        )
        cur.execute("SELECT Current_amount FROM Account WHERE Account_id = %s", (str(account_id),))
        new_balance = cur.fetchone()[0]
        conn.commit()
        conn.close()
        return new_balance, None

    @staticmethod
    def get_stats(account_id):
        """Returns (total_sent, total_received, txn_count) for the account."""
        conn = db.get_connection()
        cur = conn.cursor()
        cur.execute("SELECT COALESCE(SUM(Amount),0) FROM Transactions WHERE From_Account=? AND From_Account != 'DEPOSIT'", (str(account_id),))
        total_sent = cur.fetchone()[0]
        cur.execute("SELECT COALESCE(SUM(Amount),0) FROM Transactions WHERE To_Account=?", (str(account_id),))
        total_received = cur.fetchone()[0]
        cur.execute("SELECT COUNT(*) FROM Transactions WHERE From_Account=? OR To_Account=?", (str(account_id), str(account_id)))
        txn_count = cur.fetchone()[0]
        conn.close()
        return total_sent, total_received, txn_count

    @staticmethod
    def apply_loan(account_id, amount, months, interest_rate=15.0):
        """Create a loan record. Returns (loan_id, monthly_payment, error)."""
        loan_id = str(generate_transaction_id()) + str(random.randint(10,99))
        monthly_rate = interest_rate / 100 / 12
        if monthly_rate > 0:
            monthly_payment = amount * monthly_rate / (1 - (1 + monthly_rate) ** -months)
        else:
            monthly_payment = amount / months
        monthly_payment = round(monthly_payment, 2)
        start_date = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        conn = db.get_connection()
        cur = conn.cursor()
        cur.execute('''
            INSERT INTO Loans(Loan_ID, Account_ID, Amount, Interest_Rate, Months,
                              Monthly_Payment, Start_Date, Status, Amount_Repaid)
            VALUES(%s,%s,%s,%s,%s,%s,%s,%s,%s)
        ''', (loan_id, str(account_id), amount, interest_rate, months,
              monthly_payment, start_date, 'Active', 0.0))
        conn.commit()
        conn.close()
        return loan_id, monthly_payment, None

    @staticmethod
    def get_loans(account_id):
        conn = db.get_connection()
        cur = conn.cursor()
        cur.execute("SELECT * FROM Loans WHERE Account_ID=? ORDER BY Start_Date DESC", (str(account_id),))
        rows = cur.fetchall()
        conn.close()
        return rows

    @staticmethod
    def create_fixed_deposit(account_id, amount, days):
        """Deduct amount from account and create FD. Returns (fd_id, maturity_date, expected_return, error)."""
        conn = db.get_connection()
        cur = conn.cursor()
        cur.execute("SELECT Current_amount FROM Account WHERE Account_id=?", (str(account_id),))
        row = cur.fetchone()
        if row is None:
            conn.close()
            return None, None, None, "Account not found."
        if float(row[0]) < amount:
            conn.close()
            return None, None, None, "Insufficient balance for this fixed deposit."
        rate_map = {30: 5.0, 60: 7.0, 90: 10.0}
        interest_rate = rate_map.get(days, 5.0)
        expected_return = round(amount + amount * (interest_rate / 100) * (days / 365), 2)
        fd_id = "FD" + str(generate_transaction_id())
        start = datetime.now()
        from datetime import timedelta
        maturity = (start + timedelta(days=days)).strftime("%Y-%m-%d")
        start_str = start.strftime("%Y-%m-%d %H:%M:%S")
        cur.execute("UPDATE Account SET Current_amount = Current_amount - %s WHERE Account_id=?",
                    (amount, str(account_id)))
        cur.execute('''
            INSERT INTO FixedDeposits(FD_ID, Account_ID, Amount, Interest_Rate, Days,
                                      Start_Date, Maturity_Date, Expected_Return, Status)
            VALUES(%s,%s,%s,%s,%s,%s,%s,%s,%s)
        ''', (fd_id, str(account_id), amount, interest_rate, days, start_str, maturity, expected_return, 'Active'))
        conn.commit()
        conn.close()
        return fd_id, maturity, expected_return, None

    @staticmethod
    def get_fixed_deposits(account_id):
        conn = db.get_connection()
        cur = conn.cursor()
        cur.execute("SELECT * FROM FixedDeposits WHERE Account_ID=? ORDER BY Start_Date DESC", (str(account_id),))
        rows = cur.fetchall()
        conn.close()
        return rows

    @staticmethod
    def record_bill_payment(payment_id, account_id, biller, amount):
        """Deduct amount and record bill payment. Returns (new_balance, error)."""
        conn = db.get_connection()
        cur = conn.cursor()
        cur.execute("SELECT Current_amount FROM Account WHERE Account_id=?", (str(account_id),))
        row = cur.fetchone()
        if row is None:
            conn.close()
            return None, "Account not found."
        if float(row[0]) < amount:
            conn.close()
            return None, "Insufficient balance."
        now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        cur.execute("UPDATE Account SET Current_amount = Current_amount - %s WHERE Account_id=?",
                    (amount, str(account_id)))
        cur.execute('''
            INSERT INTO BillPayments(Payment_ID, Account_ID, Biller, Amount, Payment_Date, Status)
            VALUES(%s,%s,%s,%s,%s,%s)
        ''', (payment_id, str(account_id), biller, amount, now, 'Paid'))
        cur.execute("SELECT Current_amount FROM Account WHERE Account_id=?", (str(account_id),))
        new_balance = cur.fetchone()[0]
        conn.commit()
        conn.close()
        return new_balance, None

    @staticmethod
    def check_fraud(account_id, amount):
        """Returns list of fraud flags for a proposed transfer."""
        flags = []
        if amount > 5000:
            flags.append(f"Large transfer: GHS {amount:.2f} (threshold GHS 5,000)")
        five_mins_ago = (datetime.now() - timedelta(minutes=5)).strftime("%Y-%m-%d %H:%M:%S")
        conn = db.get_connection()
        cur = conn.cursor()
        cur.execute(
            "SELECT COUNT(*) FROM Transactions WHERE From_Account=? AND Transaction_Date > %s",
            (str(account_id), five_mins_ago)
        )
        recent_count = cur.fetchone()[0]
        conn.close()
        if recent_count >= 3:
            flags.append(f"Rapid activity: {recent_count} transfers in the last 5 minutes")
        return flags

    @staticmethod
    def create_recurring(account_id, recipient_id, amount, frequency, description):
        """Create a standing order. Returns schedule_id."""
        schedule_id = "REC" + str(generate_transaction_id())
        days = 7 if frequency == "Weekly" else 30
        next_run = (datetime.now() + timedelta(days=days)).strftime("%Y-%m-%d")
        conn = db.get_connection()
        cur = conn.cursor()
        cur.execute('''
            INSERT INTO RecurringPayments(Schedule_ID, Account_ID, Recipient_ID,
                                          Amount, Frequency, Next_Run_Date, Status, Description)
            VALUES(%s,%s,%s,%s,%s,%s,%s,%s)
        ''', (schedule_id, str(account_id), str(recipient_id), amount,
              frequency, next_run, 'Active', description))
        conn.commit()
        conn.close()
        return schedule_id

    @staticmethod
    def get_recurring(account_id):
        conn = db.get_connection()
        cur = conn.cursor()
        cur.execute(
            "SELECT * FROM RecurringPayments WHERE Account_ID=? ORDER BY Next_Run_Date ASC",
            (str(account_id),)
        )
        rows = cur.fetchall()
        conn.close()
        return rows

    @staticmethod
    def cancel_recurring(schedule_id):
        conn = db.get_connection()
        cur = conn.cursor()
        cur.execute("UPDATE RecurringPayments SET Status='Cancelled' WHERE Schedule_ID=?",
                    (schedule_id,))
        conn.commit()
        conn.close()

    @staticmethod
    def process_due_recurring():
        """Execute all active recurring payments due today or earlier."""
        today = datetime.now().strftime("%Y-%m-%d")
        conn = db.get_connection()
        cur = conn.cursor()
        cur.execute(
            "SELECT * FROM RecurringPayments WHERE Status='Active' AND Next_Run_Date <= %s",
            (today,)
        )
        due = cur.fetchall()
        conn.close()
        executed = []
        for rec in due:
            schedule_id, account_id, recipient_id, amount, frequency, _, status, description = rec
            new_balance, err = DB.transfer(account_id, recipient_id, amount)
            if not err:
                txn_id = generate_transaction_id()
                DB.record_transaction(txn_id, account_id, recipient_id, amount)
                days = 7 if frequency == "Weekly" else 30
                next_run = (datetime.now() + timedelta(days=days)).strftime("%Y-%m-%d")
                conn2 = db.get_connection()
                cur2 = conn2.cursor()
                cur2.execute("UPDATE RecurringPayments SET Next_Run_Date=? WHERE Schedule_ID=?",
                             (next_run, schedule_id))
                conn2.commit()
                conn2.close()
                executed.append((account_id, recipient_id, amount, description))
        return executed

    @staticmethod
    def add_notification(account_id, message, notif_type="info"):
        notif_id = "N" + str(generate_transaction_id())
        now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        conn = db.get_connection()
        cur = conn.cursor()
        cur.execute('''
            INSERT INTO Notifications(Notif_ID, Account_ID, Message, Type, Created_Date, Is_Read)
            VALUES(%s,%s,%s,%s,%s,%s)
        ''', (notif_id, str(account_id), message, notif_type, now, 0))
        conn.commit()
        conn.close()

    @staticmethod
    def get_notifications(account_id, limit=20):
        conn = db.get_connection()
        cur = conn.cursor()
        cur.execute(
            "SELECT * FROM Notifications WHERE Account_ID=? ORDER BY Created_Date DESC LIMIT %s",
            (str(account_id), limit)
        )
        rows = cur.fetchall()
        conn.close()
        return rows

    @staticmethod
    def mark_all_read(account_id):
        conn = db.get_connection()
        cur = conn.cursor()
        cur.execute("UPDATE Notifications SET Is_Read=1 WHERE Account_ID=?", (str(account_id),))
        conn.commit()
        conn.close()

    @staticmethod
    def get_unread_count(account_id):
        conn = db.get_connection()
        cur = conn.cursor()
        cur.execute(
            "SELECT COUNT(*) FROM Notifications WHERE Account_ID=? AND Is_Read=0",
            (str(account_id),)
        )
        count = cur.fetchone()[0]
        conn.close()
        return count

    @staticmethod
    def get_all_accounts():
        conn = db.get_connection()
        cur = conn.cursor()
        cur.execute('''
            SELECT cs.Customer_Name, cs.Customer_Email, a.Account_id, a.Current_amount,
                   CASE WHEN fa.Account_ID IS NOT NULL THEN 'Frozen' ELSE 'Active' END
            FROM Customer_services cs
            JOIN Account a ON cs.Customer_ID = a.Customer_ID
            LEFT JOIN FrozenAccounts fa ON a.Account_id = fa.Account_ID
            ORDER BY cs.Customer_Name
        ''')
        rows = cur.fetchall()
        conn.close()
        return rows

    @staticmethod
    def get_all_transactions():
        conn = db.get_connection()
        cur = conn.cursor()
        cur.execute("SELECT * FROM Transactions ORDER BY Transaction_Date DESC")
        rows = cur.fetchall()
        conn.close()
        return rows

    @staticmethod
    def freeze_account(account_id, reason="Admin action"):
        now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        conn = db.get_connection()
        cur = conn.cursor()
        cur.execute(
            "INSERT INTO FrozenAccounts(Account_ID, Frozen_Date, Reason) VALUES(%s,%s,%s) ON CONFLICT (Account_ID) DO UPDATE SET Frozen_Date=EXCLUDED.Frozen_Date, Reason=EXCLUDED.Reason",
            (str(account_id), now, reason)
        )
        conn.commit()
        conn.close()

    @staticmethod
    def unfreeze_account(account_id):
        conn = db.get_connection()
        cur = conn.cursor()
        cur.execute("DELETE FROM FrozenAccounts WHERE Account_ID=?", (str(account_id),))
        conn.commit()
        conn.close()

    @staticmethod
    def is_frozen(account_id):
        conn = db.get_connection()
        cur = conn.cursor()
        cur.execute("SELECT 1 FROM FrozenAccounts WHERE Account_ID=?", (str(account_id),))
        result = cur.fetchone()
        conn.close()
        return result is not None

    @staticmethod
    def get_bank_totals():
        """Returns (total_accounts, total_balance, total_transactions) across all accounts."""
        conn = db.get_connection()
        cur = conn.cursor()
        cur.execute("SELECT COUNT(*), COALESCE(SUM(Current_amount),0) FROM Account")
        accounts, total_bal = cur.fetchone()
        cur.execute("SELECT COUNT(*) FROM Transactions")
        total_txn = cur.fetchone()[0]
        conn.close()
        return accounts, total_bal, total_txn


# ═══════════════════════════════════════════════════════════════════════════════
#  EMAIL SERVICE  (from Bank_account.py Changepassword + send_email)
# ═══════════════════════════════════════════════════════════════════════════════
class EmailService:
    @staticmethod
    def _send(to_email, subject, body):
        msg = f"Subject: {subject}\nFrom: {BANK_EMAIL}\nTo: {to_email}\nReply-To: no-reply@gmail.com\n\n{body}"
        try:
            smtObject = smtplib.SMTP_SSL("smtp.gmail.com", 465)
            smtObject.ehlo()
            smtObject.login(BANK_EMAIL, BANK_EMAIL_PASSWORD)
            smtObject.sendmail(BANK_EMAIL, to_email, msg.encode("utf-8"))
            smtObject.quit()
            return True
        except smtplib.SMTPAuthenticationError:
            print("ERROR: Email authentication failed. Check credentials / app password.")
            return False
        except smtplib.SMTPException as e:
            print(f"ERROR: Email sending failed. {e}")
            return False
        except Exception as e:
            print(f"ERROR: {e}")
            return False

    @staticmethod
    def send_otp(to_email, otp_code):
        """Mirrors Changepassword.confirm_email() from Bank_account.py."""
        body = (
            f"One-Time Password Reset Code — Action Required\n\n"
            f"We received a request to reset your password.\n"
            f"Reset Code: {otp_code}\n\n"
            f"This code expires in 10 minutes and is for single use only.\n"
            f"Never share it with anyone.\n\n"
            f"If you did not request this, please ignore this message.\n"
            f"Your password remains unchanged."
        )
        return EmailService._send(to_email, "Access Bank — Password Reset OTP", body)

    @staticmethod
    def send_transaction_confirmation(to_email, amount, recipient, txn_id, remaining):
        """Email confirmation after a successful transfer."""
        now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        body = (
            f"Transaction Confirmation — Access Bank\n\n"
            f"Amount Sent      : GHS {amount:.2f}\n"
            f"To Account       : {recipient}\n"
            f"Transaction ID   : {txn_id}\n"
            f"Date & Time      : {now}\n"
            f"Remaining Balance: GHS {remaining:.2f}\n\n"
            f"Payment Method   : Bank Transfer\n\n"
            f"If you did not initiate this transaction, contact support immediately."
        )
        EmailService._send(to_email, "Transaction Confirmation — Access Bank", body)

    @staticmethod
    def send_welcome(to_email, name, account_id, bank_username):
        """Welcome email on successful registration (mirrors send_email.send())."""
        body = (
            f"Welcome to Access Bank, {name}!\n\n"
            f"Your account has been successfully created.\n\n"
            f"Bank Username  : {bank_username}\n"
            f"Account Number : {account_id}\n\n"
            f"Please keep your credentials safe.\n\n"
            f"Thank you for choosing Access Bank."
        )
        EmailService._send(to_email, "Welcome to Access Bank", body)

    @staticmethod
    def send_deposit_otp(to_email, otp_code, amount):
        """OTP authorization email sent before a deposit is processed."""
        body = (
            f"Deposit Authorization — Access Bank\n\n"
            f"A deposit of GHS {amount:.2f} has been requested on your account.\n\n"
            f"Authorization Code: {otp_code}\n\n"
            f"This code expires in 10 minutes and is for single use only.\n"
            f"Never share it with anyone.\n\n"
            f"If you did not initiate this deposit, please contact support immediately."
        )
        return EmailService._send(to_email, "Access Bank — Deposit Authorization OTP", body)

    @staticmethod
    def send_deposit_confirmation(to_email, amount, txn_id, new_balance):
        """Confirmation email after a successful deposit."""
        now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        body = (
            f"Deposit Confirmation — Access Bank\n\n"
            f"Amount Deposited : GHS {amount:.2f}\n"
            f"Transaction ID   : {txn_id}\n"
            f"Date & Time      : {now}\n"
            f"New Balance      : GHS {new_balance:.2f}\n\n"
            f"Your deposit has been processed successfully.\n\n"
            f"If you did not initiate this transaction, contact support immediately."
        )
        EmailService._send(to_email, "Deposit Confirmation — Access Bank", body)

    @staticmethod
    def send_bill_payment_otp(to_email, otp_code, biller, amount):
        body = (
            f"Bill Payment Authorization — Access Bank\n\n"
            f"A payment of GHS {amount:.2f} to {biller} has been requested.\n\n"
            f"Authorization Code: {otp_code}\n\n"
            f"This code expires in 10 minutes. Never share it with anyone.\n"
            f"If you did not initiate this, contact support immediately."
        )
        return EmailService._send(to_email, "Access Bank — Bill Payment OTP", body)

    @staticmethod
    def send_bill_payment_confirmation(to_email, biller, amount, payment_id, new_balance):
        now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        body = (
            f"Bill Payment Confirmation — Access Bank\n\n"
            f"Biller           : {biller}\n"
            f"Amount Paid      : GHS {amount:.2f}\n"
            f"Payment ID       : {payment_id}\n"
            f"Date & Time      : {now}\n"
            f"Remaining Balance: GHS {new_balance:.2f}\n\n"
            f"Payment processed successfully."
        )
        EmailService._send(to_email, "Bill Payment Confirmation — Access Bank", body)

    @staticmethod
    def send_loan_confirmation(to_email, loan_id, amount, months, monthly_payment):
        now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        body = (
            f"Loan Application Approved — Access Bank\n\n"
            f"Loan ID          : {loan_id}\n"
            f"Loan Amount      : GHS {amount:.2f}\n"
            f"Repayment Period : {months} months\n"
            f"Monthly Payment  : GHS {monthly_payment:.2f}\n"
            f"Interest Rate    : 15% per annum\n"
            f"Date             : {now}\n\n"
            f"Your loan has been approved. Monthly repayments are due on the same date each month.\n"
            f"Please ensure your account is funded for automatic deductions."
        )
        EmailService._send(to_email, "Loan Approved — Access Bank", body)

    @staticmethod
    def send_fd_confirmation(to_email, fd_id, amount, days, maturity_date, expected_return):
        now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        body = (
            f"Fixed Deposit Confirmation — Access Bank\n\n"
            f"FD Reference     : {fd_id}\n"
            f"Amount Locked    : GHS {amount:.2f}\n"
            f"Duration         : {days} days\n"
            f"Maturity Date    : {maturity_date}\n"
            f"Expected Return  : GHS {expected_return:.2f}\n"
            f"Date Created     : {now}\n\n"
            f"Your fixed deposit is now active. Funds will be available on the maturity date."
        )
        EmailService._send(to_email, "Fixed Deposit Confirmation — Access Bank", body)

    @staticmethod
    def send_fraud_alert(to_email, account_id, amount, flags):
        now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        flag_text = "\n".join(f"  - {f}" for f in flags)
        body = (
            f"SECURITY ALERT — Access Bank\n\n"
            f"Suspicious activity was detected on your account.\n\n"
            f"Account     : {account_id}\n"
            f"Amount      : GHS {amount:.2f}\n"
            f"Date & Time : {now}\n\n"
            f"Flags Detected:\n{flag_text}\n\n"
            f"If you did NOT authorize this transaction, contact support immediately.\n"
            f"If you authorized it, no further action is required."
        )
        EmailService._send(to_email, "SECURITY ALERT — Access Bank", body)

    @staticmethod
    def send_recurring_confirmation(to_email, schedule_id, recipient, amount, frequency, next_run):
        body = (
            f"Standing Order Created — Access Bank\n\n"
            f"Schedule ID  : {schedule_id}\n"
            f"Recipient    : {recipient}\n"
            f"Amount       : GHS {amount:.2f}\n"
            f"Frequency    : {frequency}\n"
            f"First Run    : {next_run}\n\n"
            f"Your standing order is now active.\n"
            f"To cancel it, visit the Recurring Payments section in your dashboard."
        )
        EmailService._send(to_email, "Standing Order Confirmed — Access Bank", body)


# ═══════════════════════════════════════════════════════════════════════════════
#  UTILITY FUNCTIONS  (from all source files)
# ═══════════════════════════════════════════════════════════════════════════════
def generate_account_id():
    """15-digit account number. Mirrors Register_Identity._generate_account_id()."""
    digits = list(range(1, 10))
    full = random.sample(digits, 9) + random.sample(digits, 6)
    return int("".join(str(d) for d in full))

def generate_bank_username(full_name):
    """
    Derive bank username from name (mirrors Bank_interface.py logic).
    e.g. "John Henry Doe" → "jhdjohn"
    """
    parts = full_name.strip().split()
    initials = "".join(p[0].lower() for p in parts)
    return initials + parts[0].lower()

def generate_otp():
    """6-digit OTP. Mirrors Bank_Account_main.py codegenerater()."""
    digits = list(range(1, 10))
    return "".join(str(d) for d in random.sample(digits, 6))

def generate_strong_password():
    """Mirrors Passwordgenerator.generate() from Bank_bulid_1.py."""
    WORDs = random.sample(string.ascii_uppercase, k=3)
    wrd   = random.sample(string.ascii_lowercase, k=3)
    nums  = random.sample("1234567890", k=2)
    syms  = random.sample("!£$%^&*()_+}{%s/", k=2)
    combined = WORDs + wrd + syms + nums
    random.shuffle(combined)
    return "".join(combined)

def generate_transaction_id():
    """Mirrors Transaction.transaction_generator() from Bank_account.py."""
    sample = random.sample(range(1, 10), 9)
    random.shuffle(sample)
    return int("".join(str(c) for c in sample))

def save_customer_message(msg: str):
    """
    Append customer care message to docx report.
    Mirrors customer_message() from Bank_account.py.
    """
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx is not installed. Run: pip install python-docx")
    filename = "Bank_customer_report.docx"
    if not os.path.exists(filename):
        doc = _docx.Document()
        para = doc.add_paragraph("-------------------Report----------------")
        para.add_run("\n" + msg)
        doc.save(filename)
    else:
        doc = _docx.Document(filename)
        para = doc.add_paragraph("-------------------Report----------------")
        para.add_run(msg)
        doc.save(filename)


# ═══════════════════════════════════════════════════════════════════════════════
#  SHARED STYLES
# ═══════════════════════════════════════════════════════════════════════════════
SIDEBAR_STYLE = """
QWidget {
    background-color: #2c3e50;
    color: white;
}
QPushButton {
    background-color: #34495e;
    color: white;
    border: none;
    padding: 10px 18px;
    text-align: left;
    font-size: 13px;
    border-radius: 4px;
    font-weight: bold;
}
QPushButton:hover  { background-color: #007bff; color: #ffffff; }
QPushButton:checked { background-color: #007bff; color: #ffffff; font-weight: bold; }
"""

APP_STYLE = """
QMainWindow, QWidget { background-color: #f5f5f5; color: #333333; }
QLineEdit, QTextEdit, QSpinBox, QDoubleSpinBox, QComboBox {
    border: 1px solid #ddd;
    border-radius: 4px;
    padding: 8px;
    background-color: #ffffff;
    font-size: 13px;
    color: #333333;
}
QLineEdit:focus, QTextEdit:focus, QSpinBox:focus, QDoubleSpinBox:focus {
    border: 2px solid #007bff;
    background-color: white;
}
QGroupBox {
    border: 1px solid #ddd;
    border-radius: 4px;
    padding: 10px;
    margin-top: 10px;
    background-color: #ffffff;
    font-weight: bold;
    color: #555555;
}
QGroupBox::title {
    subcontrol-origin: margin;
    subcontrol-position: top left;
    padding: 0 3px;
}
QTableWidget {
    background-color: #ffffff;
    border: 1px solid #ddd;
    border-radius: 4px;
    font-size: 12px;
}
QTableWidget::item          { padding: 5px; }
QTableWidget::item:selected { background-color: #cce5ff; color: #333333; }
QHeaderView::section {
    background-color: #f8f9fa;
    border: none;
    border-bottom: 1px solid #ddd;
    padding: 8px;
    font-weight: bold;
    color: #555555;
}
QTabWidget::pane { border: 1px solid #ddd; }
QTabBar::tab {
    background-color: #e9ecef;
    padding: 8px 20px;
    border: 1px solid #ddd;
}
QTabBar::tab:selected { background-color: white; border-bottom: 2px solid #007bff; }
"""

def _btn(text, color="#007bff", hover="#0056b3", height=40):
    b = QPushButton(text)
    b.setMinimumHeight(height)
    b.setStyleSheet(f"""
        QPushButton {{
            background-color: {color};
            color: white;
            border: none;
            border-radius: 6px;
            padding: 10px 20px;
            font-weight: bold;
            font-size: 13px;
        }}
        QPushButton:hover   {{ background-color: {hover}; }}
        QPushButton:pressed {{ background-color: #003d82; }}
    """)
    return b

def primary_btn(text, height=40):  return _btn(text, "#007bff", "#0056b3", height)
def success_btn(text, height=40):  return _btn(text, "#28a745", "#218838", height)
def warning_btn(text, height=40):  return _btn(text, "#ffc107", "#e0a800", height)
def danger_btn(text, height=40):   return _btn(text, "#dc3545", "#c82333", height)
def purple_btn(text, height=40):   return _btn(text, "#6f42c1", "#5a32a3", height)


# ═══════════════════════════════════════════════════════════════════════════════
#  OTP DIALOG
# ═══════════════════════════════════════════════════════════════════════════════
class OTPDialog(QDialog):
    """
    Modal OTP verification dialog.
    Mirrors the OTP flow from Bank_Account_main.py (com == 2) and
    Bank_account.py Changepassword.confirm_email().
    """
    def __init__(self, email, otp_code, parent=None):
        super().__init__(parent)
        self.email    = email
        self.otp_code = otp_code
        self.verified = False
        self.setWindowTitle("OTP Verification — Access Bank")
        self.setModal(True)
        self.resize(430, 300)
        self.setStyleSheet(APP_STYLE)
        self._build()

    def _build(self):
        layout = QVBoxLayout()
        layout.setContentsMargins(28, 24, 28, 24)
        layout.setSpacing(14)

        heading = QLabel("Verify Your Identity")
        heading.setFont(QFont("Arial", 15, QFont.Bold))
        layout.addWidget(heading)

        info = QLabel(
            f"A 6-digit One-Time Password has been sent to:\n"
            f"{self.email}\n\n"
            "Enter the code below. It expires in 10 minutes.\n"
            "If you didn't receive it, click Resend."
        )
        info.setWordWrap(True)
        info.setStyleSheet("color: #4a5568; font-size: 12px;")
        layout.addWidget(info)

        self.code_input = QLineEdit()
        self.code_input.setPlaceholderText("Enter OTP code")
        self.code_input.setMaxLength(10)
        self.code_input.setAlignment(Qt.AlignCenter)
        self.code_input.setStyleSheet(
            "font-size: 20px; letter-spacing: 6px; padding: 10px; border: 2px solid #3182ce;"
        )
        layout.addWidget(self.code_input)

        row = QHBoxLayout()
        resend = QPushButton("Resend Code")
        resend.setStyleSheet(
            "color:#3182ce; border:none; background:transparent; font-size:12px; padding:4px;"
        )
        resend.clicked.connect(self._resend)
        verify = primary_btn("Verify & Continue", 40)
        verify.clicked.connect(self._verify)
        row.addWidget(resend)
        row.addStretch()
        row.addWidget(verify)
        layout.addLayout(row)
        self.setLayout(layout)

    def _verify(self):
        if self.code_input.text().strip() == self.otp_code:
            self.verified = True
            self.accept()
        else:
            QMessageBox.warning(self, "Wrong Code",
                "The code you entered is incorrect.\nPlease try again or click Resend.")

    def _resend(self):
        EmailService.send_otp(self.email, self.otp_code)
        QMessageBox.information(self, "Code Resent",
            "A new OTP has been sent to your email.\n"
            "If you don't see it, check your spam folder.")


# ═══════════════════════════════════════════════════════════════════════════════
#  AUTH WINDOW  (Login + Register — from Bank_interface.py + Bank_bulid_1.py)
# ═══════════════════════════════════════════════════════════════════════════════
class AuthWindow(QWidget):
    """
    Combined login / registration window.
    Carries all validations from Bank_interface.py and Bank_bulid_1.py.
    """
    login_success = pyqtSignal(str)   # emits the logged-in email
    admin_login   = pyqtSignal()      # emits when admin credentials used

    def __init__(self):
        super().__init__()
        self.setWindowTitle("Access Bank — Welcome")
        self.setGeometry(100, 100, 540, 680)
        self.setStyleSheet(APP_STYLE)
        self._build()

    def _build(self):
        root = QVBoxLayout()
        root.setContentsMargins(44, 30, 44, 30)
        root.setSpacing(4)

        # Bank header (mirrors Bank_bulid_1.greet())
        title = QLabel("ACCESS BANK")
        title.setFont(QFont("Arial", 28, QFont.Bold))
        title.setAlignment(Qt.AlignCenter)
        title.setStyleSheet("color: #007bff; letter-spacing: 4px;")
        root.addWidget(title)

        sub = QLabel("Secure Banking Solutions")
        sub.setAlignment(Qt.AlignCenter)
        sub.setStyleSheet("color: #666666; font-size: 12px; margin-bottom: 16px;")
        root.addWidget(sub)

        tabs = QTabWidget()
        tabs.addTab(self._login_tab(),    "Sign In")
        tabs.addTab(self._register_tab(), "Create Account")
        root.addWidget(tabs)
        self.setLayout(root)

    # ── Login tab ──────────────────────────────────────────────────────────────
    def _login_tab(self):
        w = QWidget()
        layout = QVBoxLayout()
        layout.setSpacing(10)
        layout.setContentsMargins(10, 20, 10, 10)

        layout.addWidget(QLabel("Email Address"))
        self.login_email = QLineEdit()
        self.login_email.setPlaceholderText("yourname@example.com")
        layout.addWidget(self.login_email)

        layout.addWidget(QLabel("Password"))
        self.login_password = QLineEdit()
        self.login_password.setPlaceholderText("Enter your password")
        self.login_password.setEchoMode(QLineEdit.Password)
        layout.addWidget(self.login_password)

        btn = primary_btn("Sign In", 46)
        btn.clicked.connect(self._handle_login)
        layout.addSpacing(10)
        layout.addWidget(btn)
        layout.addStretch()
        w.setLayout(layout)
        return w

    # ── Register tab ───────────────────────────────────────────────────────────
    def _register_tab(self):
        w = QWidget()
        layout = QVBoxLayout()
        layout.setSpacing(8)
        layout.setContentsMargins(10, 12, 10, 10)

        def row(label, widget):
            layout.addWidget(QLabel(label))
            layout.addWidget(widget)

        self.reg_name = QLineEdit()
        self.reg_name.setPlaceholderText("e.g. John Henry Doe  (each word 3–8 letters)")
        row("Full Name *", self.reg_name)

        self.reg_age = QSpinBox()
        self.reg_age.setRange(1, 119)
        row("Age *", self.reg_age)

        self.reg_gender = QComboBox()
        self.reg_gender.addItems(["Male", "Female"])
        row("Gender *", self.reg_gender)

        self.reg_status = QComboBox()
        self.reg_status.addItems(["Single", "Married", "Student"])
        row("Status *", self.reg_status)

        self.reg_location = QLineEdit()
        self.reg_location.setPlaceholderText("Town/City (letters only)")
        row("Location *", self.reg_location)

        self.reg_phone = QLineEdit()
        self.reg_phone.setPlaceholderText("050-000-0000")
        row("Phone *", self.reg_phone)

        self.reg_email = QLineEdit()
        self.reg_email.setPlaceholderText("yourname@example.com")
        row("Email *", self.reg_email)

        # Password + Generate button
        layout.addWidget(QLabel("Password *  (8–20 chars, letters/digits/special)"))
        pwd_row = QHBoxLayout()
        self.reg_password = QLineEdit()
        self.reg_password.setPlaceholderText("Create a strong password")
        self.reg_password.setEchoMode(QLineEdit.Password)
        gen_btn = QPushButton("Generate")
        gen_btn.setFixedWidth(86)
        gen_btn.setStyleSheet(
            "background:#e2e8f0; border-radius:6px; padding:8px; font-size:11px; color:#2d3748;"
        )
        gen_btn.clicked.connect(self._generate_password)
        pwd_row.addWidget(self.reg_password)
        pwd_row.addWidget(gen_btn)
        layout.addLayout(pwd_row)

        self.reg_confirm = QLineEdit()
        self.reg_confirm.setPlaceholderText("Re-enter password")
        self.reg_confirm.setEchoMode(QLineEdit.Password)
        row("Confirm Password *", self.reg_confirm)

        btn = primary_btn("Create Account", 46)
        btn.clicked.connect(self._handle_register)
        layout.addSpacing(8)
        layout.addWidget(btn)
        w.setLayout(layout)
        return w

    def _generate_password(self):
        """Mirrors Passwordgenerator.generate() from Bank_bulid_1.py."""
        pwd = generate_strong_password()
        self.reg_password.setText(pwd)
        self.reg_password.setEchoMode(QLineEdit.Normal)
        self.reg_confirm.setText(pwd)
        QMessageBox.information(self, "Password Generated",
            f"Your generated password:\n\n{pwd}\n\n"
            "Please copy and save it somewhere safe before continuing.")

    # ── Login handler ──────────────────────────────────────────────────────────
    def _handle_login(self):
        email    = self.login_email.text().strip()
        password = self.login_password.text().strip()

        if not email or not password:
            QMessageBox.warning(self, "Missing Fields", "Please enter your email and password.")
            return

        # Admin check first — before any email format validation
        if email == ADMIN_EMAIL and password == ADMIN_PASSWORD:
            self.login_email.clear()
            self.login_password.clear()
            self.admin_login.emit()
            return

        if not re.match(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,4}", email):
            QMessageBox.warning(self, "Invalid Email", "Please enter a valid email address.")
            return

        try:
            matched = DB.login(email, password)
        except Exception as e:
            QMessageBox.critical(self, "Connection Error",
                f"Could not connect to the database.\n"
                f"Check your internet connection and try again.\n\nDetail: {e}")
            return

        if matched:
            self.login_email.clear()
            self.login_password.clear()
            self.login_success.emit(email)
        else:
            QMessageBox.critical(self, "Login Failed",
                "Invalid email or password.\nPlease try again or create an account.")

    # ── Register handler ───────────────────────────────────────────────────────
    def _handle_register(self):
        """
        Full validation chain from Bank_interface.py:
        name words (3-8 alpha), phone format, location alpha,
        email regex, strong password (8-20 chars), confirmation,
        duplicate check via Sign_up_check equivalent.
        """
        name     = self.reg_name.text().strip()
        age      = self.reg_age.value()
        gender   = self.reg_gender.currentText()
        status   = self.reg_status.currentText()
        location = self.reg_location.text().strip()
        phone    = self.reg_phone.text().strip()
        email    = self.reg_email.text().strip()
        password = self.reg_password.text().strip()
        confirm  = self.reg_confirm.text().strip()

        if not all([name, phone, email, password, confirm, location]):
            QMessageBox.warning(self, "Missing Fields", "Please fill in all required fields (*).")
            return

        # Each word in name: 3-8 alpha chars (from Bank_interface.py)
        name_re = re.compile(r"^[a-zA-Z]{3,8}$")
        for part in name.split():
            if not name_re.match(part):
                QMessageBox.warning(self, "Invalid Name",
                    "Each word in your name must be 3–8 letters only.\n"
                    "No numbers, spaces at the end, or special characters.")
                return

        # Phone format: 050-000-0000 (from Bank_interface.py)
        if not re.match(r"^\d{3}-\d{3}-\d{4}$", phone):
            QMessageBox.warning(self, "Invalid Phone",
                "Phone number must follow the format: 050-000-0000")
            return

        # Location: letters only (from Bank_interface.py)
        if not location.isalpha():
            QMessageBox.warning(self, "Invalid Location",
                "Location must contain letters only (no spaces, numbers, or special characters).")
            return

        # Email (from Bank_interface.py)
        if not re.match(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,4}", email):
            QMessageBox.warning(self, "Invalid Email",
                "Email must contain '@' and '.' e.g. name@domain.com")
            return

        # Strong password: 8-20 chars (from Bank_interface.py)
        if not re.match(r"[a-zA-Z0-9.!£$%^&*()_+|]{8,20}$", password):
            QMessageBox.warning(self, "Weak Password",
                "Password must be 8–20 characters using letters, digits, or\n"
                "special characters: !£$%^&*()_+|")
            return

        # Confirm match
        if password != confirm:
            QMessageBox.warning(self, "Password Mismatch", "Passwords do not match.")
            return

        # Duplicate check (mirrors Sign_up_check from Bank_bulid_1.py)
        try:
            duplicate = DB.is_duplicate(email)
        except Exception as e:
            QMessageBox.critical(self, "Connection Error",
                f"Could not connect to the database.\n"
                f"Check your internet connection and try again.\n\nDetail: {e}")
            return

        if duplicate:
            QMessageBox.warning(self, "Account Exists",
                "An account with this email already exists.\nPlease sign in instead.")
            return

        try:
            account_id    = generate_account_id()
            bank_username = generate_bank_username(name)   # from Bank_interface.py logic

            # Mirrors Register_Identity.register() / register_not()
            DB.register(name, age, gender, status, location, phone, email, password, account_id)

            # Welcome email (mirrors send_email.send() from Bank_bulid_1.py)
            EmailService.send_welcome(email, name, account_id, bank_username)

            QMessageBox.information(self, "Account Created — Welcome!",
                f"Your Access Bank account is ready!\n\n"
                f"Bank Username  : {bank_username}\n"
                f"Account Number : {account_id}\n\n"
                "Please save these details. You can now sign in.")

            for field in [self.reg_name, self.reg_location, self.reg_phone,
                          self.reg_email, self.reg_password, self.reg_confirm]:
                field.clear()

        except Exception as e:
            QMessageBox.critical(self, "Registration Error",
                f"Could not create account:\n{str(e)}")


# ═══════════════════════════════════════════════════════════════════════════════
#  DASHBOARD  (post-login — from Bank_Account_main.py)
# ═══════════════════════════════════════════════════════════════════════════════
class Dashboard(QMainWindow):
    """
    Full post-login dashboard.
    Mirrors all options from Bank_Account_main.py:
      1. Send Money           → Transfer page
      2. Change Password (OTP)→ Change Password page
      3. View Transactions    → History page  (was a stub in original)
      4. Customer Care        → Customer Care page
      5. Logout               → back to AuthWindow
    Plus: Profile page with account info & password generator tool.
    """
    def __init__(self, email: str):
        super().__init__()
        self.email    = email
        self.customer = DB.get_customer(email)
        self.account  = DB.get_account(self.customer[8]) if self.customer else None
        self.setWindowTitle(f"Access Bank — {self.customer[0] if self.customer else email}")
        self.setGeometry(80, 50, 1200, 760)
        self.setStyleSheet(APP_STYLE)
        self._build()
        self._refresh_transactions()
        # Check recurring payments every 60 seconds
        self._rec_timer = QTimer(self)
        self._rec_timer.timeout.connect(self._process_recurring)
        self._rec_timer.start(60000)

    def _build(self):
        central = QWidget()
        self.setCentralWidget(central)
        root = QHBoxLayout()
        root.setContentsMargins(0, 0, 0, 0)
        root.setSpacing(0)
        self.pages = QStackedWidget()
        self.pages.addWidget(self._overview_page())       # 0
        self.pages.addWidget(self._transfer_page())       # 1
        self.pages.addWidget(self._history_page())        # 2
        self.pages.addWidget(self._change_pwd_page())     # 3
        self.pages.addWidget(self._customer_care_page())  # 4
        self.pages.addWidget(self._profile_page())        # 5
        self.pages.addWidget(self._deposit_page())        # 6
        self.pages.addWidget(self._analytics_page())     # 7
        self.pages.addWidget(self._bill_payments_page()) # 8
        self.pages.addWidget(self._loans_page())         # 9
        self.pages.addWidget(self._fixed_deposit_page()) # 10
        self.pages.addWidget(self._recurring_page())     # 11
        self.pages.addWidget(self._notifications_page()) # 12

        root.addWidget(self._sidebar())
        root.addWidget(self.pages, 1)
        central.setLayout(root)

    # ── Sidebar ────────────────────────────────────────────────────────────────
    def _sidebar(self):
        sb = QWidget()
        sb.setFixedWidth(226)
        sb.setStyleSheet(SIDEBAR_STYLE)
        layout = QVBoxLayout()
        layout.setContentsMargins(12, 0, 12, 20)
        layout.setSpacing(2)

        # Bank title (mirrors Bank_bulid_1.greet())
        title = QLabel("ACCESS BANK")
        title.setFont(QFont("Arial", 18, QFont.Bold))
        title.setStyleSheet("color: #ffffff; padding: 22px 10px 4px 10px; letter-spacing: 3px;")
        layout.addWidget(title)

        name_lbl = QLabel(self.customer[0] if self.customer else "")
        name_lbl.setStyleSheet("color: #aaaaaa; font-size: 11px; padding: 0 10px 16px 10px;")
        layout.addWidget(name_lbl)

        sep = QFrame()
        sep.setFrameShape(QFrame.HLine)
        sep.setStyleSheet("color: #3d566e; margin: 0 4px 10px 4px;")
        layout.addWidget(sep)

        nav_items = [
            ("   Overview",             0),
            ("   Send Money",           1),
            ("   Deposit Money",        6),
            ("   Analytics",            7),
            ("   Bill Payments",        8),
            ("   Loans",                9),
            ("   Fixed Deposits",      10),
            ("   Recurring Payments",  11),
            ("   Notifications",       12),
            ("   Transaction History",  2),
            ("   Change Password",      3),
            ("   Customer Care",        4),
            ("   My Profile",           5),
        ]
        self._nav_btns = []
        self._notif_btn = None
        for label, idx in nav_items:
            btn = QPushButton(label)
            btn.setCheckable(True)
            btn.setMinimumHeight(40)
            btn.clicked.connect(lambda _, i=idx: self._nav(i))
            self._nav_btns.append(btn)
            layout.addWidget(btn)
            if idx == 12:
                self._notif_btn = btn
        self._update_notif_badge()

        layout.addStretch()

        logout = QPushButton("   Logout")
        logout.setMinimumHeight(46)
        logout.setStyleSheet("""
            QPushButton {
                background-color: #dc3545;
                color: white;
                border: none;
                border-radius: 4px;
                padding: 10px;
                font-weight: bold;
                text-align: left;
            }
            QPushButton:hover { background-color: #c82333; color: white; }
        """)
        logout.clicked.connect(self._logout)
        layout.addWidget(logout)

        sb.setLayout(layout)
        self._nav(0)
        return sb

    def _nav(self, index):
        self.pages.setCurrentIndex(index)
        nav_indices = [0, 1, 6, 7, 8, 9, 10, 11, 12, 2, 3, 4, 5]
        for i, btn in enumerate(self._nav_btns):
            btn.setChecked(nav_indices[i] == index)
        if index == 0:
            self._refresh_balance()
        if index == 2:
            self._refresh_transactions()
        if index == 7:
            self._refresh_analytics()
        if index == 8:
            self._refresh_bill_table()
        if index == 9:
            self._refresh_loan_table()
        if index == 10:
            self._refresh_fd_table()
        if index == 11:
            self._refresh_recurring_table()
        if index == 12:
            self._refresh_notifications()

    # ── Overview page ──────────────────────────────────────────────────────────
    def _overview_page(self):
        page = QWidget()
        layout = QVBoxLayout()
        layout.setContentsMargins(34, 30, 34, 30)
        layout.setSpacing(22)

        name = self.customer[0] if self.customer else "Customer"
        greet = QLabel(f"Welcome back, {name}!")
        greet.setFont(QFont("Arial", 22, QFont.Bold))
        greet.setStyleSheet("color: #333333;")
        layout.addWidget(greet)

        # Balance card
        bal_card = QWidget()
        bal_card.setMinimumHeight(130)
        bal_card.setStyleSheet("""
            background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                stop:0 #007bff, stop:1 #0056b3);
            border-radius: 8px;
        """)
        bcl = QVBoxLayout()
        bcl.setContentsMargins(28, 20, 28, 20)

        lbl = QLabel("Current Balance")
        lbl.setStyleSheet("color: rgba(255,255,255,0.65); font-size: 13px;")
        bcl.addWidget(lbl)

        balance = self.account[2] if self.account and self.account[2] is not None else 0.0
        self.bal_display = QLabel(f"GHS {balance:,.2f}")
        self.bal_display.setFont(QFont("Arial", 30, QFont.Bold))
        self.bal_display.setStyleSheet("color: white;")
        bcl.addWidget(self.bal_display)

        acct_no = str(self.customer[8]) if self.customer else "N/A"
        acct_lbl = QLabel(f"Account Number: {acct_no}")
        acct_lbl.setStyleSheet("color: rgba(255,255,255,0.6); font-size: 12px;")
        bcl.addWidget(acct_lbl)

        bal_card.setLayout(bcl)
        layout.addWidget(bal_card)

        # Quick actions (mirrors Bank_Account_main.py menu)
        qa = QGroupBox("Quick Actions")
        qa_row = QHBoxLayout()
        qa_row.setSpacing(14)

        sb = primary_btn("  Send Money")
        sb.clicked.connect(lambda: self._nav(1))
        qa_row.addWidget(sb)

        db = success_btn("  Deposit Money")
        db.clicked.connect(lambda: self._nav(6))
        qa_row.addWidget(db)

        bb = _btn("  Bill Payments", "#17a2b8", "#138496")
        bb.clicked.connect(lambda: self._nav(8))
        qa_row.addWidget(bb)

        lb = _btn("  Loans", "#fd7e14", "#e8610a")
        lb.clicked.connect(lambda: self._nav(9))
        qa_row.addWidget(lb)

        hb = warning_btn("  View Transactions")
        hb.clicked.connect(lambda: self._nav(2))
        qa_row.addWidget(hb)

        cb = purple_btn("  Customer Care")
        cb.clicked.connect(lambda: self._nav(4))
        qa_row.addWidget(cb)

        qa.setLayout(qa_row)
        layout.addWidget(qa)
        layout.addStretch()
        page.setLayout(layout)
        return page

    def _refresh_balance(self):
        if not hasattr(self, 'bal_display') or not self.customer:
            return
        acc = DB.get_account(self.customer[8])
        if acc:
            self.account = acc
            self.bal_display.setText(f"GHS {acc[2]:,.2f}")

    # ── Transfer page (mirrors Bank_Account_main.py com == 1) ─────────────────
    def _transfer_page(self):
        page = QWidget()
        layout = QVBoxLayout()
        layout.setContentsMargins(34, 30, 34, 30)
        layout.setSpacing(20)

        title = QLabel("Send Money")
        title.setFont(QFont("Arial", 20, QFont.Bold))
        layout.addWidget(title)

        note = QLabel(
            "Enter the recipient's 15-digit account number and the amount to transfer.\n"
            "A confirmation email will be sent to you after the transfer."
        )
        note.setWordWrap(True)
        note.setStyleSheet("color: #718096; font-size: 12px;")
        layout.addWidget(note)

        form_card = QGroupBox("Transfer Details")
        form = QFormLayout()
        form.setSpacing(14)

        self.t_recipient = QLineEdit()
        self.t_recipient.setPlaceholderText("15-digit account number")
        form.addRow("Recipient Account:", self.t_recipient)

        self.t_amount = QDoubleSpinBox()
        self.t_amount.setRange(0.01, 9999999.99)
        self.t_amount.setDecimals(2)
        self.t_amount.setPrefix("GHS ")
        form.addRow("Amount:", self.t_amount)

        self.t_desc = QLineEdit()
        self.t_desc.setPlaceholderText("Optional description / note")
        form.addRow("Description:", self.t_desc)

        form_card.setLayout(form)
        layout.addWidget(form_card)

        send_btn = primary_btn("  Send Transfer", 50)
        send_btn.clicked.connect(self._handle_transfer)
        layout.addWidget(send_btn)
        layout.addStretch()
        page.setLayout(layout)
        return page

    def _handle_transfer(self):
        recipient = self.t_recipient.text().strip()
        amount    = self.t_amount.value()

        if not recipient:
            QMessageBox.warning(self, "Missing Field", "Please enter a recipient account number.")
            return
        if not recipient.isdigit() or len(recipient) != 15:
            QMessageBox.warning(self, "Invalid Account",
                "Account number must be exactly 15 digits.")
            return
        if amount <= 0:
            QMessageBox.warning(self, "Invalid Amount", "Amount must be greater than 0.")
            return
        sender_id = str(self.customer[8])
        if recipient == sender_id:
            QMessageBox.warning(self, "Invalid", "You cannot transfer to your own account.")
            return

        # Check if account is frozen
        if DB.is_frozen(sender_id):
            QMessageBox.critical(self, "Account Frozen",
                "Your account has been frozen. Please contact support.")
            return

        # Fraud detection
        flags = DB.check_fraud(sender_id, amount)
        if flags:
            flag_msg = "\n".join(f"  • {f}" for f in flags)
            EmailService.send_fraud_alert(self.email, sender_id, amount, flags)
            reply = QMessageBox.warning(self, "Suspicious Activity Detected",
                f"The following concerns were flagged on this transfer:\n\n{flag_msg}\n\n"
                "A security alert has been sent to your registered email.\n\n"
                "Do you still want to proceed%s",
                QMessageBox.Yes | QMessageBox.No
            )
            if reply == QMessageBox.No:
                return
            DB.add_notification(sender_id,
                f"Fraud alert: suspicious transfer of GHS {amount:.2f} was flagged but allowed.",
                "warning")
            self._update_notif_badge()

        # mirrors Transaction.check() + Transaction.send()
        new_balance, err = DB.transfer(sender_id, recipient, amount)
        if err:
            QMessageBox.warning(self, "Transfer Failed", err)
            return

        # mirrors Transaction.transaction_generator() + transaction_record() + transaction_report()
        txn_id = generate_transaction_id()
        DB.record_transaction(txn_id, sender_id, recipient, amount)

        # email confirmation
        EmailService.send_transaction_confirmation(self.email, amount, recipient, txn_id, new_balance)

        DB.add_notification(sender_id,
            f"Transfer of GHS {amount:.2f} sent to account {recipient}. New balance: GHS {new_balance:.2f}",
            "info")
        self._update_notif_badge()

        self.account = DB.get_account(self.customer[8])
        self._refresh_balance()

        now = datetime.now()
        QMessageBox.information(self, "Transfer Successful",
            f"Your transfer was successful!\n\n"
            f"Amount          : GHS {amount:.2f}\n"
            f"To Account      : {recipient}\n"
            f"Transaction ID  : {txn_id}\n"
            f"Date & Time     : {now.strftime('%Y-%m-%d %H:%M:%S')}\n"
            f"Payment Method  : Bank Transfer\n"
            f"Remaining Balance: GHS {new_balance:.2f}\n\n"
            "A confirmation has been sent to your email.\n"
            "If you did not initiate this, contact support immediately.")

        self.t_recipient.clear()
        self.t_amount.setValue(0.01)
        self.t_desc.clear()
        self._refresh_transactions()

    # ── Transaction history page (was a stub in Bank_Account_main.py) ─────────
    def _history_page(self):
        page = QWidget()
        layout = QVBoxLayout()
        layout.setContentsMargins(34, 30, 34, 30)
        layout.setSpacing(16)

        hdr = QHBoxLayout()
        title = QLabel("Transaction History")
        title.setFont(QFont("Arial", 20, QFont.Bold))
        hdr.addWidget(title)
        hdr.addStretch()
        export_btn = QPushButton("Export CSV")
        export_btn.setStyleSheet(
            "background:#28a745; color:white; border-radius:6px; padding:8px 18px; font-weight:bold;"
        )
        export_btn.clicked.connect(self._export_statement)
        hdr.addWidget(export_btn)
        ref = QPushButton("Refresh")
        ref.setStyleSheet(
            "background:#e2e8f0; border-radius:6px; padding:8px 18px; color:#2d3748;"
        )
        ref.clicked.connect(self._refresh_transactions)
        hdr.addWidget(ref)
        layout.addLayout(hdr)

        self.txn_table = QTableWidget()
        self.txn_table.setColumnCount(6)
        self.txn_table.setHorizontalHeaderLabels(
            ["Transaction ID", "From Account", "To Account", "Amount (GHS)", "Date & Time", "Status"]
        )
        self.txn_table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        self.txn_table.setEditTriggers(QTableWidget.NoEditTriggers)
        self.txn_table.setAlternatingRowColors(True)
        self.txn_table.setSelectionBehavior(QTableWidget.SelectRows)
        self.txn_table.verticalHeader().setVisible(False)
        layout.addWidget(self.txn_table)
        page.setLayout(layout)
        return page

    def _refresh_transactions(self):
        if not hasattr(self, 'txn_table') or not self.customer:
            return
        rows = DB.get_transactions(str(self.customer[8]))
        self.txn_table.setRowCount(len(rows))
        for r, row in enumerate(rows):
            for c, val in enumerate(row):
                item = QTableWidgetItem(
                    f"GHS {float(val):,.2f}" if c == 3 else str(val)
                )
                item.setTextAlignment(Qt.AlignCenter)
                if c == 3:
                    item.setForeground(QColor("#276749"))
                if c == 5 and str(val) == "Completed":
                    item.setForeground(QColor("#276749"))
                self.txn_table.setItem(r, c, item)

    def _export_statement(self):
        if not self.customer:
            return
        rows = DB.get_transactions(str(self.customer[8]))
        filename = f"statement_{self.customer[8]}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"
        with open(filename, "w", newline="") as f:
            writer = csv.writer(f)
            writer.writerow(["Transaction ID", "From Account", "To Account",
                             "Amount (GHS)", "Date & Time", "Status"])
            for row in rows:
                writer.writerow([row[0], row[1], row[2],
                                 f"{float(row[3]):.2f}", row[4], row[5]])
        QMessageBox.information(self, "Export Successful",
            f"Account statement exported to:\n{filename}\n\n"
            f"{len(rows)} transaction(s) included.")

    # ── Change password page (mirrors Bank_Account_main.py com == 2) ──────────
    def _change_pwd_page(self):
        page = QWidget()
        layout = QVBoxLayout()
        layout.setContentsMargins(34, 30, 34, 30)
        layout.setSpacing(20)

        title = QLabel("Change Password")
        title.setFont(QFont("Arial", 20, QFont.Bold))
        layout.addWidget(title)

        info = QLabel(
            "First verify your current password, then a One-Time Password will be sent\n"
            "to your registered email address to confirm the change."
        )
        info.setWordWrap(True)
        info.setStyleSheet("color: #718096; font-size: 12px;")
        layout.addWidget(info)

        pwd_card = QGroupBox("Update Password")
        form = QFormLayout()
        form.setSpacing(14)

        self.cp_current = QLineEdit()
        self.cp_current.setEchoMode(QLineEdit.Password)
        self.cp_current.setPlaceholderText("Enter your current password")
        form.addRow("Current Password:", self.cp_current)

        self.cp_new = QLineEdit()
        self.cp_new.setEchoMode(QLineEdit.Password)
        self.cp_new.setPlaceholderText("8–20 chars: letters, digits, !.{}/+-_")
        form.addRow("New Password:", self.cp_new)

        self.cp_confirm = QLineEdit()
        self.cp_confirm.setEchoMode(QLineEdit.Password)
        self.cp_confirm.setPlaceholderText("Re-enter new password")
        form.addRow("Confirm Password:", self.cp_confirm)

        pwd_card.setLayout(form)
        layout.addWidget(pwd_card)

        update_btn = primary_btn("Update Password  (sends OTP)", 50)
        update_btn.clicked.connect(self._handle_change_pwd)
        layout.addWidget(update_btn)
        layout.addStretch()
        page.setLayout(layout)
        return page

    def _handle_change_pwd(self):
        """
        Mirrors Bank_Account_main.py com==2 with OTP flow.
        Uses Changepassword.confirm_email() logic via EmailService.send_otp()
        and Changepassword.change() logic via DB.update_password().
        """
        current = self.cp_current.text()
        new     = self.cp_new.text()
        confirm = self.cp_confirm.text()

        if not all([current, new, confirm]):
            QMessageBox.warning(self, "Missing Fields", "Please fill in all password fields.")
            return
        if not DB.login(self.email, current):
            QMessageBox.warning(self, "Wrong Password", "Your current password is incorrect.")
            return
        if new != confirm:
            QMessageBox.warning(self, "Mismatch", "New passwords do not match.")
            return
        # mirrors Bank_Account_main.py regex: r"[A-Za-z0-9!.{}/+-_]{8,20}"
        if not re.match(r"[A-Za-z0-9!.{}/+\-_]{8,20}$", new):
            QMessageBox.warning(self, "Weak Password",
                "Password must be 8–20 characters using:\nletters, digits, or !.{}/+-_")
            return

        otp = generate_otp()
        sent = EmailService.send_otp(self.email, otp)   # mirrors Changepassword.confirm_email()
        if not sent:
            QMessageBox.warning(self, "Email Error",
                "Could not send OTP. Check email credentials.\n"
                "The OTP dialog will still open — but you may not receive the code.")

        dialog = OTPDialog(self.email, otp, self)
        if dialog.exec_() == QDialog.Accepted and dialog.verified:
            DB.update_password(self.email, new)          # mirrors Changepassword.change()
            QMessageBox.information(self, "Success", "Your password has been updated successfully!")
            self.cp_current.clear()
            self.cp_new.clear()
            self.cp_confirm.clear()

    # ── Customer Care page (mirrors Bank_Account_main.py com == 4) ────────────
    def _customer_care_page(self):
        page = QWidget()
        layout = QVBoxLayout()
        layout.setContentsMargins(34, 30, 34, 30)
        layout.setSpacing(20)

        title = QLabel("Customer Care")
        title.setFont(QFont("Arial", 20, QFont.Bold))
        layout.addWidget(title)

        info = QLabel(
            "Hi there! Need help with your account or have a question%s\n"
            "Just type your message below — our Customer Care team is here for you\n"
            "and will respond shortly."
        )
        info.setWordWrap(True)
        info.setStyleSheet("color: #718096; font-size: 13px;")
        layout.addWidget(info)

        msg_card = QGroupBox("Your Message")
        msg_layout = QVBoxLayout()
        msg_layout.setSpacing(12)

        self.care_input = QTextEdit()
        self.care_input.setPlaceholderText("Enter your message here...")
        self.care_input.setMinimumHeight(180)
        msg_layout.addWidget(self.care_input)

        btn_row = QHBoxLayout()
        cancel_btn = warning_btn("Cancel")
        cancel_btn.clicked.connect(self._cancel_care)
        send_btn = primary_btn("Send Message", 44)
        send_btn.clicked.connect(self._handle_care)
        btn_row.addWidget(cancel_btn)
        btn_row.addStretch()
        btn_row.addWidget(send_btn)
        msg_layout.addLayout(btn_row)

        self.care_status = QLabel("")
        self.care_status.setStyleSheet("color: #276749; font-size: 12px;")
        msg_layout.addWidget(self.care_status)

        msg_card.setLayout(msg_layout)
        layout.addWidget(msg_card)
        layout.addStretch()
        page.setLayout(layout)
        return page

    def _cancel_care(self):
        self.care_input.clear()
        self.care_status.setText("Message cancelled.")
        QTimer.singleShot(3000, lambda: self.care_status.setText(""))

    def _handle_care(self):
        """Mirrors Bank_Account_main.py com==4 + customer_message() from Bank_account.py."""
        msg = self.care_input.toPlainText().strip()
        if not msg:
            QMessageBox.warning(self, "Empty Message",
                "Please enter a message before sending.")
            return
        reply = QMessageBox.question(
            self, "Confirm",
            "Do you wish to send this message%s\n\n"
            "Our Customer Care team will attend to you shortly.",
            QMessageBox.Yes | QMessageBox.No
        )
        if reply == QMessageBox.Yes:
            try:
                save_customer_message(msg)   # mirrors customer_message()
                self.care_input.clear()
                self.care_status.setText(
                    "Message sent successfully. We will attend to you shortly."
                )
                QTimer.singleShot(6000, lambda: self.care_status.setText(""))
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to send message:\n{str(e)}")
        else:
            self.care_status.setText("Your request to send a message was cancelled.")
            QTimer.singleShot(3000, lambda: self.care_status.setText(""))

    # ── Profile page (account info + password generator) ──────────────────────
    def _profile_page(self):
        page = QWidget()
        layout = QVBoxLayout()
        layout.setContentsMargins(34, 30, 34, 30)
        layout.setSpacing(20)

        title = QLabel("My Profile")
        title.setFont(QFont("Arial", 20, QFont.Bold))
        layout.addWidget(title)

        if self.customer:
            # Account info
            info_card = QGroupBox("Account Information")
            form = QFormLayout()
            form.setSpacing(10)
            c = self.customer
            bank_uname = generate_bank_username(c[0])   # from Bank_interface.py
            fields = [
                ("Full Name",      c[0]),
                ("Bank Username",  bank_uname),
                ("Age",            c[1]),
                ("Gender",         c[2]),
                ("Status",         c[3]),
                ("Location",       c[4]),
                ("Phone",          c[5]),
                ("Email",          c[6]),
                ("Account Number", str(c[8])),
            ]
            for label, value in fields:
                lbl = QLabel(str(value))
                lbl.setStyleSheet(
                    "color: #2d3748; font-size: 13px; padding: 3px 0;"
                )
                form.addRow(f"{label}:", lbl)
            info_card.setLayout(form)
            layout.addWidget(info_card)

        # Password generator tool (mirrors Passwordgenerator.generate() from Bank_bulid_1.py)
        gen_card = QGroupBox("Password Generator Tool")
        gen_layout = QVBoxLayout()
        gen_layout.setSpacing(10)

        gen_note = QLabel(
            "Need a strong password%s Click the button to generate one.\n"
            "Copy and save it before using it."
        )
        gen_note.setStyleSheet("color: #718096; font-size: 12px;")
        gen_layout.addWidget(gen_note)

        self.gen_output = QLineEdit()
        self.gen_output.setReadOnly(True)
        self.gen_output.setPlaceholderText("Generated password will appear here")
        self.gen_output.setStyleSheet(
            "font-size: 14px; letter-spacing: 2px; padding: 10px; background: #f7fafc;"
        )
        gen_layout.addWidget(self.gen_output)

        gen_btn = success_btn("Generate Strong Password", 44)
        gen_btn.clicked.connect(
            lambda: self.gen_output.setText(generate_strong_password())
        )
        gen_layout.addWidget(gen_btn)
        gen_card.setLayout(gen_layout)
        layout.addWidget(gen_card)

        layout.addStretch()
        page.setLayout(layout)
        return page

    # ── Deposit page ──────────────────────────────────────────────────────────
    def _deposit_page(self):
        page = QWidget()
        layout = QVBoxLayout()
        layout.setContentsMargins(34, 30, 34, 30)
        layout.setSpacing(20)

        title = QLabel("Deposit Money")
        title.setFont(QFont("Arial", 20, QFont.Bold))
        layout.addWidget(title)

        note = QLabel(
            "Enter the amount you wish to deposit into your account.\n"
            "A One-Time Password (OTP) will be sent to your registered email\n"
            "to authorize the deposit before it is processed."
        )
        note.setWordWrap(True)
        note.setStyleSheet("color: #718096; font-size: 12px;")
        layout.addWidget(note)

        form_card = QGroupBox("Deposit Details")
        form = QFormLayout()
        form.setSpacing(14)

        self.dep_amount = QDoubleSpinBox()
        self.dep_amount.setRange(0.01, 9999999.99)
        self.dep_amount.setDecimals(2)
        self.dep_amount.setPrefix("GHS ")
        form.addRow("Deposit Amount:", self.dep_amount)

        form_card.setLayout(form)
        layout.addWidget(form_card)

        dep_btn = success_btn("  Deposit Funds  (sends OTP)", 50)
        dep_btn.clicked.connect(self._handle_deposit)
        layout.addWidget(dep_btn)
        layout.addStretch()
        page.setLayout(layout)
        return page

    def _handle_deposit(self):
        amount = self.dep_amount.value()
        if amount <= 0:
            QMessageBox.warning(self, "Invalid Amount", "Amount must be greater than 0.")
            return

        otp = generate_otp()
        sent = EmailService.send_deposit_otp(self.email, otp, amount)
        if not sent:
            QMessageBox.warning(self, "Email Error",
                "Could not send OTP. Check email credentials.\n"
                "The OTP dialog will still open — but you may not receive the code.")

        dialog = OTPDialog(self.email, otp, self)
        dialog.setWindowTitle("Deposit Authorization — Access Bank")
        if dialog.exec_() == QDialog.Accepted and dialog.verified:
            account_id = str(self.customer[8])
            new_balance, err = DB.deposit(account_id, amount)
            if err:
                QMessageBox.critical(self, "Deposit Failed", err)
                return

            txn_id = generate_transaction_id()
            DB.record_transaction(txn_id, "DEPOSIT", account_id, amount)
            EmailService.send_deposit_confirmation(self.email, amount, txn_id, new_balance)
            DB.add_notification(account_id,
                f"Deposit of GHS {amount:.2f} successful. New balance: GHS {new_balance:.2f}", "info")
            self._update_notif_badge()

            self.account = DB.get_account(self.customer[8])
            self._refresh_balance()
            self._refresh_transactions()

            now = datetime.now()
            QMessageBox.information(self, "Deposit Successful",
                f"Your deposit was processed successfully!\n\n"
                f"Amount Deposited : GHS {amount:.2f}\n"
                f"Transaction ID   : {txn_id}\n"
                f"Date & Time      : {now.strftime('%Y-%m-%d %H:%M:%S')}\n"
                f"New Balance      : GHS {new_balance:.2f}\n\n"
                "A confirmation has been sent to your email.")

            self.dep_amount.setValue(0.01)

    # ── Analytics page (page 7) ───────────────────────────────────────────────
    def _analytics_page(self):
        page = QWidget()
        layout = QVBoxLayout()
        layout.setContentsMargins(34, 30, 34, 30)
        layout.setSpacing(22)

        hdr = QHBoxLayout()
        title = QLabel("Analytics")
        title.setFont(QFont("Arial", 20, QFont.Bold))
        hdr.addWidget(title)
        hdr.addStretch()
        ref_btn = QPushButton("Refresh")
        ref_btn.setStyleSheet(
            "background:#e2e8f0; border-radius:6px; padding:8px 18px; color:#2d3748;"
        )
        ref_btn.clicked.connect(self._refresh_analytics)
        hdr.addWidget(ref_btn)
        layout.addLayout(hdr)

        # Stat cards row
        cards_row = QHBoxLayout()
        cards_row.setSpacing(16)

        def stat_card(title_text, subtitle_text):
            card = QGroupBox()
            card_layout = QVBoxLayout()
            card_layout.setSpacing(4)
            val_lbl = QLabel("—")
            val_lbl.setFont(QFont("Arial", 22, QFont.Bold))
            val_lbl.setStyleSheet("color: #007bff;")
            val_lbl.setAlignment(Qt.AlignCenter)
            sub_lbl = QLabel(subtitle_text)
            sub_lbl.setStyleSheet("color: #718096; font-size: 12px;")
            sub_lbl.setAlignment(Qt.AlignCenter)
            card_layout.addWidget(val_lbl)
            card_layout.addWidget(sub_lbl)
            card.setLayout(card_layout)
            return card, val_lbl

        sent_card,    self.an_sent     = stat_card("Total Sent",        "Total Sent")
        recv_card,    self.an_received = stat_card("Total Received",    "Total Received")
        count_card,   self.an_count    = stat_card("Transactions",      "Total Transactions")
        bal_card_an,  self.an_balance  = stat_card("Current Balance",   "Current Balance")

        for card in [sent_card, recv_card, count_card, bal_card_an]:
            cards_row.addWidget(card)
        layout.addLayout(cards_row)

        # Recent activity table
        recent_grp = QGroupBox("Recent Activity (Last 5 Transactions)")
        recent_layout = QVBoxLayout()
        self.an_table = QTableWidget()
        self.an_table.setColumnCount(4)
        self.an_table.setHorizontalHeaderLabels(["ID", "Type", "Amount", "Date"])
        self.an_table.setEditTriggers(QTableWidget.NoEditTriggers)
        self.an_table.setAlternatingRowColors(True)
        self.an_table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        self.an_table.verticalHeader().setVisible(False)
        recent_layout.addWidget(self.an_table)
        recent_grp.setLayout(recent_layout)
        layout.addWidget(recent_grp)

        layout.addStretch()
        page.setLayout(layout)
        return page

    def _refresh_analytics(self):
        if not hasattr(self, 'an_sent') or not self.customer:
            return
        account_id = str(self.customer[8])
        total_sent, total_received, txn_count = DB.get_stats(account_id)
        acc = DB.get_account(self.customer[8])
        balance = acc[2] if acc and acc[2] is not None else 0.0
        self.an_sent.setText(f"GHS {total_sent:,.2f}")
        self.an_received.setText(f"GHS {total_received:,.2f}")
        self.an_count.setText(str(txn_count))
        self.an_balance.setText(f"GHS {balance:,.2f}")

        rows = DB.get_transactions(account_id)[:5]
        self.an_table.setRowCount(len(rows))
        for r, row in enumerate(rows):
            txn_id  = str(row[0])
            from_ac = str(row[1])
            to_ac   = str(row[2])
            amount  = float(row[3])
            date    = str(row[4])
            if from_ac == account_id:
                txn_type = f"Sent → {to_ac}"
            else:
                txn_type = f"Received ← {from_ac}"
            vals = [txn_id, txn_type, f"GHS {amount:,.2f}", date]
            for c, val in enumerate(vals):
                item = QTableWidgetItem(val)
                item.setTextAlignment(Qt.AlignCenter)
                self.an_table.setItem(r, c, item)

    # ── Bill Payments page (page 8) ───────────────────────────────────────────
    def _bill_payments_page(self):
        page = QWidget()
        layout = QVBoxLayout()
        layout.setContentsMargins(34, 30, 34, 30)
        layout.setSpacing(20)

        title = QLabel("Bill Payments")
        title.setFont(QFont("Arial", 20, QFont.Bold))
        layout.addWidget(title)

        form_card = QGroupBox("Pay a Bill")
        form = QFormLayout()
        form.setSpacing(14)

        self.bill_biller = QComboBox()
        self.bill_biller.addItems([
            "ECG (Electricity)",
            "GWCL (Water)",
            "MTN Mobile Money",
            "Vodafone Cash",
            "AirtelTigo Money",
            "DStv Subscription",
            "StarTimes Subscription",
            "GoTV Subscription",
        ])
        form.addRow("Biller:", self.bill_biller)

        self.bill_amount = QDoubleSpinBox()
        self.bill_amount.setRange(1.0, 5000.0)
        self.bill_amount.setDecimals(2)
        self.bill_amount.setPrefix("GHS ")
        form.addRow("Amount:", self.bill_amount)

        form_card.setLayout(form)
        layout.addWidget(form_card)

        pay_btn = success_btn("  Pay Bill (sends OTP)", 50)
        pay_btn.clicked.connect(self._handle_bill_payment)
        layout.addWidget(pay_btn)

        history_grp = QGroupBox("Payment History")
        history_layout = QVBoxLayout()
        self.bill_table = QTableWidget()
        self.bill_table.setColumnCount(4)
        self.bill_table.setHorizontalHeaderLabels(["Payment ID", "Biller", "Amount", "Date"])
        self.bill_table.setEditTriggers(QTableWidget.NoEditTriggers)
        self.bill_table.setAlternatingRowColors(True)
        self.bill_table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        self.bill_table.verticalHeader().setVisible(False)
        history_layout.addWidget(self.bill_table)
        history_grp.setLayout(history_layout)
        layout.addWidget(history_grp)

        layout.addStretch()
        page.setLayout(layout)
        return page

    def _handle_bill_payment(self):
        biller = self.bill_biller.currentText()
        amount = self.bill_amount.value()
        account_id = str(self.customer[8])

        otp = generate_otp()
        sent = EmailService.send_bill_payment_otp(self.email, otp, biller, amount)
        if not sent:
            QMessageBox.warning(self, "Email Error",
                "Could not send OTP. Check email credentials.\n"
                "The OTP dialog will still open — but you may not receive the code.")

        dialog = OTPDialog(self.email, otp, self)
        dialog.setWindowTitle("Bill Payment Authorization — Access Bank")
        if dialog.exec_() == QDialog.Accepted and dialog.verified:
            payment_id = str(generate_transaction_id())
            new_balance, err = DB.record_bill_payment(payment_id, account_id, biller, amount)
            if err:
                QMessageBox.critical(self, "Payment Failed", err)
                return

            txn_id = generate_transaction_id()
            DB.record_transaction(txn_id, account_id, f"BILLS-{biller[:3].upper()}", amount)
            EmailService.send_bill_payment_confirmation(self.email, biller, amount, payment_id, new_balance)
            DB.add_notification(account_id,
                f"Bill payment of GHS {amount:.2f} to {biller} completed.", "info")
            self._update_notif_badge()

            self.account = DB.get_account(self.customer[8])
            self._refresh_balance()
            self._refresh_bill_table()

            QMessageBox.information(self, "Payment Successful",
                f"Bill payment processed successfully!\n\n"
                f"Biller          : {biller}\n"
                f"Amount Paid     : GHS {amount:.2f}\n"
                f"Payment ID      : {payment_id}\n"
                f"Remaining Balance: GHS {new_balance:.2f}\n\n"
                "A confirmation has been sent to your email.")

            self.bill_amount.setValue(1.0)

    def _refresh_bill_table(self):
        if not hasattr(self, 'bill_table') or not self.customer:
            return
        conn = db.get_connection()
        cur = conn.cursor()
        cur.execute(
            "SELECT * FROM BillPayments WHERE Account_ID=? ORDER BY Payment_Date DESC",
            (str(self.customer[8]),)
        )
        rows = cur.fetchall()
        conn.close()
        self.bill_table.setRowCount(len(rows))
        for r, row in enumerate(rows):
            payment_id = str(row[0])
            biller     = str(row[2])
            amount     = f"GHS {float(row[3]):,.2f}"
            date       = str(row[4])
            for c, val in enumerate([payment_id, biller, amount, date]):
                item = QTableWidgetItem(val)
                item.setTextAlignment(Qt.AlignCenter)
                self.bill_table.setItem(r, c, item)

    # ── Loans page (page 9) ───────────────────────────────────────────────────
    def _loans_page(self):
        page = QWidget()
        layout = QVBoxLayout()
        layout.setContentsMargins(34, 30, 34, 30)
        layout.setSpacing(20)

        title = QLabel("Loan Management")
        title.setFont(QFont("Arial", 20, QFont.Bold))
        layout.addWidget(title)

        form_card = QGroupBox("Apply for a Loan")
        form = QFormLayout()
        form.setSpacing(14)

        self.loan_amount = QDoubleSpinBox()
        self.loan_amount.setRange(100.0, 50000.0)
        self.loan_amount.setDecimals(2)
        self.loan_amount.setPrefix("GHS ")
        form.addRow("Loan Amount:", self.loan_amount)

        self.loan_months = QComboBox()
        self.loan_months.addItems(["6 months", "12 months", "24 months", "36 months"])
        form.addRow("Repayment Period:", self.loan_months)

        rate_lbl = QLabel("Interest Rate: 15% per annum")
        rate_lbl.setStyleSheet("color: #718096; font-size: 12px;")
        form.addRow("", rate_lbl)

        form_card.setLayout(form)
        layout.addWidget(form_card)

        apply_btn = primary_btn("  Apply for Loan", 50)
        apply_btn.clicked.connect(self._handle_loan)
        layout.addWidget(apply_btn)

        history_grp = QGroupBox("My Loans")
        history_layout = QVBoxLayout()
        self.loan_table = QTableWidget()
        self.loan_table.setColumnCount(6)
        self.loan_table.setHorizontalHeaderLabels(
            ["Loan ID", "Amount", "Months", "Monthly Payment", "Start Date", "Status"]
        )
        self.loan_table.setEditTriggers(QTableWidget.NoEditTriggers)
        self.loan_table.setAlternatingRowColors(True)
        self.loan_table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        self.loan_table.verticalHeader().setVisible(False)
        history_layout.addWidget(self.loan_table)
        history_grp.setLayout(history_layout)
        layout.addWidget(history_grp)

        layout.addStretch()
        page.setLayout(layout)
        return page

    def _handle_loan(self):
        amount = self.loan_amount.value()
        months_text = self.loan_months.currentText()
        months = int(months_text.split()[0])
        account_id = str(self.customer[8])

        # Calculate preview monthly payment
        monthly_rate = 15.0 / 100 / 12
        monthly_payment_preview = round(
            amount * monthly_rate / (1 - (1 + monthly_rate) ** -months), 2
        )

        reply = QMessageBox.question(
            self, "Confirm Loan Application",
            f"Apply for GHS {amount:.2f} over {months} months at 15% p.a.%s\n"
            f"Monthly Payment: GHS {monthly_payment_preview:.2f}\n\n"
            "Proceed with this loan application%s",
            QMessageBox.Yes | QMessageBox.No
        )
        if reply != QMessageBox.Yes:
            return

        loan_id, monthly_payment, err = DB.apply_loan(account_id, amount, months)
        if err:
            QMessageBox.critical(self, "Loan Failed", err)
            return

        new_balance, dep_err = DB.deposit(account_id, amount)
        if dep_err:
            QMessageBox.critical(self, "Credit Failed", dep_err)
            return

        txn_id = generate_transaction_id()
        DB.record_transaction(txn_id, "LOAN-CREDIT", account_id, amount)
        EmailService.send_loan_confirmation(self.email, loan_id, amount, months, monthly_payment)
        DB.add_notification(account_id,
            f"Loan of GHS {amount:.2f} approved over {months} months. Monthly: GHS {monthly_payment:.2f}",
            "info")
        self._update_notif_badge()

        self.account = DB.get_account(self.customer[8])
        self._refresh_balance()
        self._refresh_loan_table()

        QMessageBox.information(self, "Loan Approved",
            f"Your loan has been approved!\n\n"
            f"Loan ID         : {loan_id}\n"
            f"Amount          : GHS {amount:.2f}\n"
            f"Repayment       : {months} months\n"
            f"Monthly Payment : GHS {monthly_payment:.2f}\n\n"
            "Funds have been credited to your account.\n"
            "A confirmation has been sent to your email.")

        self.loan_amount.setValue(100.0)

    def _refresh_loan_table(self):
        if not hasattr(self, 'loan_table') or not self.customer:
            return
        rows = DB.get_loans(str(self.customer[8]))
        self.loan_table.setRowCount(len(rows))
        for r, row in enumerate(rows):
            loan_id    = str(row[0])
            amount     = f"GHS {float(row[2]):,.2f}"
            months     = str(row[4])
            monthly    = f"GHS {float(row[5]):,.2f}"
            start_date = str(row[6])
            status     = str(row[7])
            for c, val in enumerate([loan_id, amount, months, monthly, start_date, status]):
                item = QTableWidgetItem(val)
                item.setTextAlignment(Qt.AlignCenter)
                self.loan_table.setItem(r, c, item)

    # ── Fixed Deposits page (page 10) ─────────────────────────────────────────
    def _fixed_deposit_page(self):
        page = QWidget()
        layout = QVBoxLayout()
        layout.setContentsMargins(34, 30, 34, 30)
        layout.setSpacing(20)

        title = QLabel("Fixed Deposits")
        title.setFont(QFont("Arial", 20, QFont.Bold))
        layout.addWidget(title)

        form_card = QGroupBox("Create a Fixed Deposit")
        form = QFormLayout()
        form.setSpacing(14)

        self.fd_amount = QDoubleSpinBox()
        self.fd_amount.setRange(100.0, 500000.0)
        self.fd_amount.setDecimals(2)
        self.fd_amount.setPrefix("GHS ")
        form.addRow("Amount:", self.fd_amount)

        self.fd_days = QComboBox()
        self.fd_days.addItems([
            "30 days — 5% p.a.",
            "60 days — 7% p.a.",
            "90 days — 10% p.a.",
        ])
        form.addRow("Duration:", self.fd_days)

        info_lbl = QLabel("Funds are locked for the selected period. Early withdrawal is not available.")
        info_lbl.setStyleSheet("color: #718096; font-size: 12px;")
        info_lbl.setWordWrap(True)
        form.addRow("", info_lbl)

        form_card.setLayout(form)
        layout.addWidget(form_card)

        create_btn = warning_btn("  Create Fixed Deposit", 50)
        create_btn.clicked.connect(self._handle_fd)
        layout.addWidget(create_btn)

        history_grp = QGroupBox("My Fixed Deposits")
        history_layout = QVBoxLayout()
        self.fd_table = QTableWidget()
        self.fd_table.setColumnCount(6)
        self.fd_table.setHorizontalHeaderLabels(
            ["FD ID", "Amount", "Days", "Rate", "Maturity Date", "Expected Return"]
        )
        self.fd_table.setEditTriggers(QTableWidget.NoEditTriggers)
        self.fd_table.setAlternatingRowColors(True)
        self.fd_table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        self.fd_table.verticalHeader().setVisible(False)
        history_layout.addWidget(self.fd_table)
        history_grp.setLayout(history_layout)
        layout.addWidget(history_grp)

        layout.addStretch()
        page.setLayout(layout)
        return page

    def _handle_fd(self):
        amount = self.fd_amount.value()
        days_text = self.fd_days.currentText()
        days = int(days_text[:2].strip())
        account_id = str(self.customer[8])

        rate_map = {30: 5.0, 60: 7.0, 90: 10.0}
        interest_rate = rate_map.get(days, 5.0)
        expected_return_preview = round(amount + amount * (interest_rate / 100) * (days / 365), 2)
        from datetime import timedelta
        maturity_preview = (datetime.now() + timedelta(days=days)).strftime("%Y-%m-%d")

        reply = QMessageBox.question(
            self, "Confirm Fixed Deposit",
            f"Amount          : GHS {amount:.2f}\n"
            f"Duration        : {days} days\n"
            f"Interest Rate   : {interest_rate}% p.a.\n"
            f"Maturity Date   : {maturity_preview}\n"
            f"Expected Return : GHS {expected_return_preview:.2f}\n\n"
            "Funds will be locked until maturity. Proceed%s",
            QMessageBox.Yes | QMessageBox.No
        )
        if reply != QMessageBox.Yes:
            return

        fd_id, maturity_date, expected_return, err = DB.create_fixed_deposit(account_id, amount, days)
        if err:
            QMessageBox.critical(self, "Fixed Deposit Failed", err)
            return

        txn_id = generate_transaction_id()
        DB.record_transaction(txn_id, account_id, "FD-LOCK", amount)
        EmailService.send_fd_confirmation(self.email, fd_id, amount, days, maturity_date, expected_return)
        DB.add_notification(account_id,
            f"Fixed deposit of GHS {amount:.2f} created. Matures {maturity_date}. Return: GHS {expected_return:.2f}",
            "info")
        self._update_notif_badge()

        self.account = DB.get_account(self.customer[8])
        self._refresh_balance()
        self._refresh_fd_table()

        QMessageBox.information(self, "Fixed Deposit Created",
            f"Your fixed deposit has been created!\n\n"
            f"FD Reference    : {fd_id}\n"
            f"Amount Locked   : GHS {amount:.2f}\n"
            f"Duration        : {days} days\n"
            f"Maturity Date   : {maturity_date}\n"
            f"Expected Return : GHS {expected_return:.2f}\n\n"
            "A confirmation has been sent to your email.")

        self.fd_amount.setValue(100.0)

    def _refresh_fd_table(self):
        if not hasattr(self, 'fd_table') or not self.customer:
            return
        rows = DB.get_fixed_deposits(str(self.customer[8]))
        self.fd_table.setRowCount(len(rows))
        for r, row in enumerate(rows):
            fd_id    = str(row[0])
            amount   = f"GHS {float(row[2]):,.2f}"
            days     = str(row[4])
            rate     = f"{float(row[3])}%"
            maturity = str(row[6])
            exp_ret  = f"GHS {float(row[7]):,.2f}"
            for c, val in enumerate([fd_id, amount, days, rate, maturity, exp_ret]):
                item = QTableWidgetItem(val)
                item.setTextAlignment(Qt.AlignCenter)
                self.fd_table.setItem(r, c, item)

    # ── Recurring Payments page (page 11) ────────────────────────────────────
    def _recurring_page(self):
        page = QWidget()
        layout = QVBoxLayout()
        layout.setContentsMargins(34, 30, 34, 30)
        layout.setSpacing(20)

        title = QLabel("Recurring Payments")
        title.setFont(QFont("Arial", 20, QFont.Bold))
        layout.addWidget(title)

        note = QLabel(
            "Set up automatic standing orders to transfer money to another account\n"
            "on a weekly or monthly schedule."
        )
        note.setWordWrap(True)
        note.setStyleSheet("color: #718096; font-size: 12px;")
        layout.addWidget(note)

        form_card = QGroupBox("New Standing Order")
        form = QFormLayout()
        form.setSpacing(14)

        self.rec_recipient = QLineEdit()
        self.rec_recipient.setPlaceholderText("15-digit account number")
        form.addRow("Recipient Account:", self.rec_recipient)

        self.rec_amount = QDoubleSpinBox()
        self.rec_amount.setRange(1.0, 50000.0)
        self.rec_amount.setDecimals(2)
        self.rec_amount.setPrefix("GHS ")
        form.addRow("Amount:", self.rec_amount)

        self.rec_frequency = QComboBox()
        self.rec_frequency.addItems(["Weekly", "Monthly"])
        form.addRow("Frequency:", self.rec_frequency)

        self.rec_desc = QLineEdit()
        self.rec_desc.setPlaceholderText("e.g. Rent, Savings, Family support")
        form.addRow("Description:", self.rec_desc)

        form_card.setLayout(form)
        layout.addWidget(form_card)

        setup_btn = primary_btn("  Set Up Standing Order", 50)
        setup_btn.clicked.connect(self._handle_recurring)
        layout.addWidget(setup_btn)

        history_grp = QGroupBox("My Standing Orders")
        history_layout = QVBoxLayout()
        self.rec_table = QTableWidget()
        self.rec_table.setColumnCount(6)
        self.rec_table.setHorizontalHeaderLabels(
            ["Schedule ID", "Recipient", "Amount", "Frequency", "Next Run", "Status"]
        )
        self.rec_table.setEditTriggers(QTableWidget.NoEditTriggers)
        self.rec_table.setAlternatingRowColors(True)
        self.rec_table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        self.rec_table.verticalHeader().setVisible(False)
        self.rec_table.setSelectionBehavior(QTableWidget.SelectRows)
        history_layout.addWidget(self.rec_table)

        cancel_btn = danger_btn("Cancel Selected Order", 36)
        cancel_btn.clicked.connect(self._cancel_selected_recurring)
        history_layout.addWidget(cancel_btn)

        history_grp.setLayout(history_layout)
        layout.addWidget(history_grp)

        layout.addStretch()
        page.setLayout(layout)
        return page

    def _handle_recurring(self):
        recipient = self.rec_recipient.text().strip()
        amount    = self.rec_amount.value()
        frequency = self.rec_frequency.currentText()
        desc      = self.rec_desc.text().strip() or "Standing Order"
        account_id = str(self.customer[8])

        if not recipient or not recipient.isdigit() or len(recipient) != 15:
            QMessageBox.warning(self, "Invalid Account",
                "Recipient account must be exactly 15 digits.")
            return
        if recipient == account_id:
            QMessageBox.warning(self, "Invalid", "Cannot set up a standing order to your own account.")
            return

        days = 7 if frequency == "Weekly" else 30
        from datetime import timedelta
        next_run = (datetime.now() + timedelta(days=days)).strftime("%Y-%m-%d")

        reply = QMessageBox.question(self, "Confirm Standing Order",
            f"Recipient : {recipient}\n"
            f"Amount    : GHS {amount:.2f}\n"
            f"Frequency : {frequency}\n"
            f"First Run : {next_run}\n\n"
            "Set up this standing order%s",
            QMessageBox.Yes | QMessageBox.No
        )
        if reply != QMessageBox.Yes:
            return

        schedule_id = DB.create_recurring(account_id, recipient, amount, frequency, desc)
        EmailService.send_recurring_confirmation(self.email, schedule_id, recipient, amount, frequency, next_run)
        DB.add_notification(account_id,
            f"Standing order set up: GHS {amount:.2f} to {recipient} every {frequency.lower()}.", "info")
        self._update_notif_badge()
        self._refresh_recurring_table()

        QMessageBox.information(self, "Standing Order Created",
            f"Your standing order has been set up!\n\n"
            f"Schedule ID : {schedule_id}\n"
            f"Recipient   : {recipient}\n"
            f"Amount      : GHS {amount:.2f}\n"
            f"Frequency   : {frequency}\n"
            f"First Run   : {next_run}\n\n"
            "A confirmation has been sent to your email.")

        self.rec_recipient.clear()
        self.rec_amount.setValue(1.0)
        self.rec_desc.clear()

    def _cancel_selected_recurring(self):
        row = self.rec_table.currentRow()
        if row < 0:
            QMessageBox.warning(self, "No Selection", "Please select a standing order to cancel.")
            return
        schedule_id = self.rec_table.item(row, 0).text()
        status      = self.rec_table.item(row, 5).text()
        if status == "Cancelled":
            QMessageBox.information(self, "Already Cancelled", "This standing order is already cancelled.")
            return
        reply = QMessageBox.question(self, "Confirm Cancellation",
            f"Cancel standing order {schedule_id}%s\nThis cannot be undone.",
            QMessageBox.Yes | QMessageBox.No
        )
        if reply == QMessageBox.Yes:
            DB.cancel_recurring(schedule_id)
            self._refresh_recurring_table()
            QMessageBox.information(self, "Cancelled", "Standing order has been cancelled.")

    def _refresh_recurring_table(self):
        if not hasattr(self, 'rec_table') or not self.customer:
            return
        rows = DB.get_recurring(str(self.customer[8]))
        self.rec_table.setRowCount(len(rows))
        for r, row in enumerate(rows):
            vals = [str(row[0]), str(row[2]), f"GHS {float(row[3]):,.2f}",
                    str(row[4]), str(row[5]), str(row[6])]
            for c, val in enumerate(vals):
                item = QTableWidgetItem(val)
                item.setTextAlignment(Qt.AlignCenter)
                if c == 5 and val == "Cancelled":
                    item.setForeground(QColor("#dc3545"))
                elif c == 5 and val == "Active":
                    item.setForeground(QColor("#276749"))
                self.rec_table.setItem(r, c, item)

    def _process_recurring(self):
        """Called every 60s by QTimer to execute due standing orders."""
        executed = DB.process_due_recurring()
        for account_id, recipient, amount, desc in executed:
            DB.add_notification(account_id,
                f"Standing order executed: GHS {amount:.2f} sent to {recipient} ({desc}).", "info")
        if executed and self.customer and str(self.customer[8]) in [e[0] for e in executed]:
            self._refresh_balance()
            self._update_notif_badge()

    # ── Notifications page (page 12) ──────────────────────────────────────────
    def _notifications_page(self):
        page = QWidget()
        layout = QVBoxLayout()
        layout.setContentsMargins(34, 30, 34, 30)
        layout.setSpacing(16)

        hdr = QHBoxLayout()
        title = QLabel("Notifications")
        title.setFont(QFont("Arial", 20, QFont.Bold))
        hdr.addWidget(title)
        hdr.addStretch()
        mark_btn = QPushButton("Mark All Read")
        mark_btn.setStyleSheet(
            "background:#6c757d; color:white; border-radius:6px; padding:8px 18px; font-weight:bold;"
        )
        mark_btn.clicked.connect(self._mark_all_read)
        hdr.addWidget(mark_btn)
        layout.addLayout(hdr)

        self.notif_table = QTableWidget()
        self.notif_table.setColumnCount(4)
        self.notif_table.setHorizontalHeaderLabels(["#", "Message", "Date", "Read"])
        self.notif_table.setEditTriggers(QTableWidget.NoEditTriggers)
        self.notif_table.setAlternatingRowColors(True)
        self.notif_table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        self.notif_table.verticalHeader().setVisible(False)
        self.notif_table.setSelectionBehavior(QTableWidget.SelectRows)
        layout.addWidget(self.notif_table)
        page.setLayout(layout)
        return page

    def _refresh_notifications(self):
        if not hasattr(self, 'notif_table') or not self.customer:
            return
        rows = DB.get_notifications(str(self.customer[8]))
        self.notif_table.setRowCount(len(rows))
        for r, row in enumerate(rows):
            is_read = bool(row[5])
            vals = [str(r + 1), str(row[2]), str(row[4]), "Yes" if is_read else "Unread"]
            for c, val in enumerate(vals):
                item = QTableWidgetItem(val)
                item.setTextAlignment(Qt.AlignCenter if c != 1 else Qt.AlignLeft | Qt.AlignVCenter)
                if not is_read:
                    item.setForeground(QColor("#007bff"))
                    item.setFont(QFont("Arial", 10, QFont.Bold))
                self.notif_table.setItem(r, c, item)

    def _mark_all_read(self):
        if not self.customer:
            return
        DB.mark_all_read(str(self.customer[8]))
        self._refresh_notifications()
        self._update_notif_badge()

    def _update_notif_badge(self):
        if not hasattr(self, '_notif_btn') or self._notif_btn is None or not self.customer:
            return
        count = DB.get_unread_count(str(self.customer[8]))
        if count > 0:
            self._notif_btn.setText(f"   Notifications  ({count})")
        else:
            self._notif_btn.setText("   Notifications")

    # ── Logout (mirrors Bank_Account_main.py com == 5) ────────────────────────
    def _logout(self):
        self.close()
        app = QApplication.instance()
        if hasattr(app, 'auth_window'):
            app.auth_window.show()


# ═══════════════════════════════════════════════════════════════════════════════
#  ADMIN DASHBOARD
# ═══════════════════════════════════════════════════════════════════════════════
class AdminDashboard(QMainWindow):
    """Separate admin view — accessed via ADMIN_EMAIL / ADMIN_PASSWORD credentials."""

    def __init__(self):
        super().__init__()
        self.setWindowTitle("Access Bank — Admin Portal")
        self.setGeometry(60, 40, 1240, 800)
        self.setStyleSheet(APP_STYLE)
        self._build()

    def _build(self):
        central = QWidget()
        self.setCentralWidget(central)
        root = QHBoxLayout()
        root.setContentsMargins(0, 0, 0, 0)
        root.setSpacing(0)

        self.pages = QStackedWidget()
        self.pages.addWidget(self._admin_overview())   # 0
        self.pages.addWidget(self._admin_accounts())   # 1
        self.pages.addWidget(self._admin_transactions()) # 2

        root.addWidget(self._sidebar())
        root.addWidget(self.pages, 1)
        central.setLayout(root)
        self._refresh_overview()

    def _sidebar(self):
        sb = QWidget()
        sb.setFixedWidth(220)
        sb.setStyleSheet(SIDEBAR_STYLE)
        layout = QVBoxLayout()
        layout.setContentsMargins(12, 0, 12, 20)
        layout.setSpacing(2)

        title = QLabel("ADMIN PORTAL")
        title.setFont(QFont("Arial", 16, QFont.Bold))
        title.setStyleSheet("color: #e74c3c; padding: 22px 10px 4px 10px; letter-spacing: 2px;")
        layout.addWidget(title)

        sub = QLabel("Access Bank")
        sub.setStyleSheet("color: #aaaaaa; font-size: 11px; padding: 0 10px 16px 10px;")
        layout.addWidget(sub)

        sep = QFrame()
        sep.setFrameShape(QFrame.HLine)
        sep.setStyleSheet("color: #3d566e; margin: 0 4px 10px 4px;")
        layout.addWidget(sep)

        self._admin_btns = []
        for label, idx in [("   Overview", 0), ("   All Accounts", 1), ("   All Transactions", 2)]:
            btn = QPushButton(label)
            btn.setCheckable(True)
            btn.setMinimumHeight(46)
            btn.clicked.connect(lambda _, i=idx: self._admin_nav(i))
            self._admin_btns.append(btn)
            layout.addWidget(btn)

        layout.addStretch()

        logout = QPushButton("   Logout")
        logout.setMinimumHeight(46)
        logout.setStyleSheet("""
            QPushButton {
                background-color: #dc3545; color: white; border: none;
                border-radius: 4px; padding: 10px; font-weight: bold; text-align: left;
            }
            QPushButton:hover { background-color: #c82333; }
        """)
        logout.clicked.connect(self._logout)
        layout.addWidget(logout)
        sb.setLayout(layout)
        self._admin_nav(0)
        return sb

    def _admin_nav(self, index):
        self.pages.setCurrentIndex(index)
        for i, btn in enumerate(self._admin_btns):
            btn.setChecked(i == index)
        if index == 0:
            self._refresh_overview()
        if index == 1:
            self._refresh_accounts()
        if index == 2:
            self._refresh_all_txns()

    def _admin_overview(self):
        page = QWidget()
        layout = QVBoxLayout()
        layout.setContentsMargins(34, 30, 34, 30)
        layout.setSpacing(22)

        title = QLabel("Bank Overview")
        title.setFont(QFont("Arial", 22, QFont.Bold))
        layout.addWidget(title)

        cards_row = QHBoxLayout()
        cards_row.setSpacing(16)

        def stat_card(subtitle):
            card = QGroupBox()
            cl = QVBoxLayout()
            cl.setSpacing(4)
            val = QLabel("—")
            val.setFont(QFont("Arial", 24, QFont.Bold))
            val.setStyleSheet("color: #e74c3c;")
            val.setAlignment(Qt.AlignCenter)
            sub = QLabel(subtitle)
            sub.setStyleSheet("color: #718096; font-size: 12px;")
            sub.setAlignment(Qt.AlignCenter)
            cl.addWidget(val)
            cl.addWidget(sub)
            card.setLayout(cl)
            return card, val

        ac_card, self.adm_accounts  = stat_card("Total Accounts")
        bl_card, self.adm_balance   = stat_card("Total Bank Balance")
        tx_card, self.adm_txns      = stat_card("Total Transactions")

        for c in [ac_card, bl_card, tx_card]:
            cards_row.addWidget(c)
        layout.addLayout(cards_row)

        ref_btn = QPushButton("Refresh Stats")
        ref_btn.setStyleSheet(
            "background:#e2e8f0; border-radius:6px; padding:8px 18px; color:#2d3748;"
        )
        ref_btn.clicked.connect(self._refresh_overview)
        layout.addWidget(ref_btn, alignment=Qt.AlignLeft)
        layout.addStretch()
        page.setLayout(layout)
        return page

    def _refresh_overview(self):
        if not hasattr(self, 'adm_accounts'):
            return
        accounts, total_bal, total_txn = DB.get_bank_totals()
        self.adm_accounts.setText(str(accounts))
        self.adm_balance.setText(f"GHS {total_bal:,.2f}")
        self.adm_txns.setText(str(total_txn))

    def _admin_accounts(self):
        page = QWidget()
        layout = QVBoxLayout()
        layout.setContentsMargins(34, 30, 34, 30)
        layout.setSpacing(16)

        hdr = QHBoxLayout()
        title = QLabel("All Customer Accounts")
        title.setFont(QFont("Arial", 20, QFont.Bold))
        hdr.addWidget(title)
        hdr.addStretch()
        ref = QPushButton("Refresh")
        ref.setStyleSheet("background:#e2e8f0; border-radius:6px; padding:8px 18px; color:#2d3748;")
        ref.clicked.connect(self._refresh_accounts)
        hdr.addWidget(ref)
        layout.addLayout(hdr)

        self.adm_acc_table = QTableWidget()
        self.adm_acc_table.setColumnCount(5)
        self.adm_acc_table.setHorizontalHeaderLabels(
            ["Customer Name", "Email", "Account Number", "Balance (GHS)", "Status"]
        )
        self.adm_acc_table.setEditTriggers(QTableWidget.NoEditTriggers)
        self.adm_acc_table.setAlternatingRowColors(True)
        self.adm_acc_table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        self.adm_acc_table.verticalHeader().setVisible(False)
        self.adm_acc_table.setSelectionBehavior(QTableWidget.SelectRows)
        layout.addWidget(self.adm_acc_table)

        btn_row = QHBoxLayout()
        freeze_btn = danger_btn("Freeze Selected Account", 38)
        freeze_btn.clicked.connect(self._freeze_selected)
        unfreeze_btn = success_btn("Unfreeze Selected Account", 38)
        unfreeze_btn.clicked.connect(self._unfreeze_selected)
        btn_row.addWidget(freeze_btn)
        btn_row.addWidget(unfreeze_btn)
        btn_row.addStretch()
        layout.addLayout(btn_row)

        page.setLayout(layout)
        return page

    def _refresh_accounts(self):
        if not hasattr(self, 'adm_acc_table'):
            return
        rows = DB.get_all_accounts()
        self.adm_acc_table.setRowCount(len(rows))
        for r, row in enumerate(rows):
            vals = [str(row[0]), str(row[1]), str(row[2]),
                    f"{float(row[3]):,.2f}", str(row[4])]
            for c, val in enumerate(vals):
                item = QTableWidgetItem(val)
                item.setTextAlignment(Qt.AlignCenter)
                if c == 4 and val == "Frozen":
                    item.setForeground(QColor("#dc3545"))
                elif c == 4 and val == "Active":
                    item.setForeground(QColor("#276749"))
                self.adm_acc_table.setItem(r, c, item)

    def _freeze_selected(self):
        row = self.adm_acc_table.currentRow()
        if row < 0:
            QMessageBox.warning(self, "No Selection", "Please select an account first.")
            return
        account_id = self.adm_acc_table.item(row, 2).text()
        name       = self.adm_acc_table.item(row, 0).text()
        reply = QMessageBox.question(self, "Confirm Freeze",
            f"Freeze account {account_id} ({name})%s\n"
            "The customer will not be able to make transfers.",
            QMessageBox.Yes | QMessageBox.No
        )
        if reply == QMessageBox.Yes:
            DB.freeze_account(account_id)
            self._refresh_accounts()
            QMessageBox.information(self, "Account Frozen", f"Account {account_id} has been frozen.")

    def _unfreeze_selected(self):
        row = self.adm_acc_table.currentRow()
        if row < 0:
            QMessageBox.warning(self, "No Selection", "Please select an account first.")
            return
        account_id = self.adm_acc_table.item(row, 2).text()
        name       = self.adm_acc_table.item(row, 0).text()
        reply = QMessageBox.question(self, "Confirm Unfreeze",
            f"Unfreeze account {account_id} ({name})%s",
            QMessageBox.Yes | QMessageBox.No
        )
        if reply == QMessageBox.Yes:
            DB.unfreeze_account(account_id)
            self._refresh_accounts()
            QMessageBox.information(self, "Account Unfrozen", f"Account {account_id} has been unfrozen.")

    def _admin_transactions(self):
        page = QWidget()
        layout = QVBoxLayout()
        layout.setContentsMargins(34, 30, 34, 30)
        layout.setSpacing(16)

        hdr = QHBoxLayout()
        title = QLabel("All Transactions")
        title.setFont(QFont("Arial", 20, QFont.Bold))
        hdr.addWidget(title)
        hdr.addStretch()
        ref = QPushButton("Refresh")
        ref.setStyleSheet("background:#e2e8f0; border-radius:6px; padding:8px 18px; color:#2d3748;")
        ref.clicked.connect(self._refresh_all_txns)
        hdr.addWidget(ref)
        layout.addLayout(hdr)

        self.adm_txn_table = QTableWidget()
        self.adm_txn_table.setColumnCount(6)
        self.adm_txn_table.setHorizontalHeaderLabels(
            ["Transaction ID", "From Account", "To Account", "Amount (GHS)", "Date & Time", "Status"]
        )
        self.adm_txn_table.setEditTriggers(QTableWidget.NoEditTriggers)
        self.adm_txn_table.setAlternatingRowColors(True)
        self.adm_txn_table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        self.adm_txn_table.verticalHeader().setVisible(False)
        layout.addWidget(self.adm_txn_table)
        page.setLayout(layout)
        return page

    def _refresh_all_txns(self):
        if not hasattr(self, 'adm_txn_table'):
            return
        rows = DB.get_all_transactions()
        self.adm_txn_table.setRowCount(len(rows))
        for r, row in enumerate(rows):
            for c, val in enumerate(row):
                item = QTableWidgetItem(
                    f"GHS {float(val):,.2f}" if c == 3 else str(val)
                )
                item.setTextAlignment(Qt.AlignCenter)
                if c == 5 and str(val) == "Completed":
                    item.setForeground(QColor("#276749"))
                self.adm_txn_table.setItem(r, c, item)

    def _logout(self):
        self.close()
        app = QApplication.instance()
        if hasattr(app, 'auth_window'):
            app.auth_window.show()


# ═══════════════════════════════════════════════════════════════════════════════
#  APPLICATION ENTRY POINT
# ═══════════════════════════════════════════════════════════════════════════════
class AccessBankApp(QApplication):
    def __init__(self):
        super().__init__(sys.argv)
        self.setApplicationName("Access Bank")
        DB.init()
        self.auth_window = AuthWindow()
        self.auth_window.login_success.connect(self._on_login)
        self.auth_window.admin_login.connect(self._on_admin_login)
        self.auth_window.show()

    def _on_login(self, email: str):
        self.auth_window.hide()
        self.dashboard = Dashboard(email)
        self.dashboard.show()

    def _on_admin_login(self):
        self.auth_window.hide()
        self.admin_dashboard = AdminDashboard()
        self.admin_dashboard.show()


if __name__ == "__main__":
    app = AccessBankApp()
    sys.exit(app.exec_())
